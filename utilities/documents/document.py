# from .guides import (UsageOfDynamicSystemsGuide, Guide, EngeneeringDrawingGuide, DevelopmentGuide, IntroToPandasGuide,
#                     BasicsOfODESystemGuide, BasicsOfDynSysImplementationGuide, BasicsOfReportingGuide, ResearchProjectGuidelines,
#                     InterimProjectGuidelines,IntroDynPyProjectGuidelines, BasicsOfReportComponentImplementationGuide,
#                     GithubSynchroGuide, IntroToCocalcGuide)
# from sympy import *
import datetime
import os
import shutil
from typing import List, Optional, Union

from pylatex import (  # Section, Subsection, Subsubsection, Itemize,  HorizontalSpace, Description, Marker
    Command,
    Document,
    NewPage,
    Package,
    Tabularx,
)
from pylatex.base_classes import Environment

# from pylatex.section import Paragraph, Chapter
from pylatex.utils import NoEscape  # italic,

from ..report import (
    CurrentContainer,
    IPMarkdown,
    Markdown,
    ObjectCode,
    ReportText,
    display,
)

imports_str = """
#Create file output
#Create file Images
#In file output create bibliography as .bib file (biblio.bib)

'''# File content begin


@misc{DynPi,
author={GitHub},
title="bogumilchilinski/dynpy",
url={https://github.com/bogumilchilinski/dynpy}",
note="Accessed:2024-05-04"
}


@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
'''# File content end

from dynpy.utilities.report import *
from dynpy.utilities.documents.document import WutThesis
#from dynpy.models.odes.linear import SpringMassDamperMSM
from sympy import symbols, Eq
from sympy.printing.latex import latex
from IPython.display import display, Math

from dynpy.utilities.adaptable import TimeDataFrame
import pandas as pd

global_width=('8cm')#setting picture default size
global_height=('6cm')
doc = WutThesis('./output/thesis_name')
#doc.preamble.append(NoEscape(r'\\usepackage[MeX]{polski}')) #to set polish as main language"""

thesis_introduction_str = """
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_intro = Section('Section that presents text reporting')
CurrentContainer(sec_intro)

sub_problem_outline = Subsection('Outline of the problem')
CurrentContainer(sub_problem_outline)
display(ReportText('This subsection provides information about investigated problem. '*100))

sub_obj_assum = Subsection('Objectives and assumptions')
CurrentContainer(sub_obj_assum)
display(ReportText('This subsection provides objectives and assumptions. '*100))

sub_SOT = Subsection('State of the art')
CurrentContainer(sub_SOT)
display(ReportText('This subsection provides state of the art with multiple reference for \\cite{DynPi}. '*50))
display(Markdown('Remeber about `biblio.bib` file with referencens. Place it in working directory and subfolder with output files. '*50))


sub_methodology = Subsection('Methodology')
CurrentContainer(sub_methodology)
display(ReportText('This subsection provides methodology. '*100))"""

thesis_introduction_str_research = """
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_intro = Section('Section that presents text reporting')
CurrentContainer(sec_intro)

sub_problem_outline = Subsection('Outline of the problem')
CurrentContainer(sub_problem_outline)
display(ReportText('This subsection provides information about investigated problem. '*100))

sub_obj_assum = Subsection('Objectives and assumptions')
CurrentContainer(sub_obj_assum)
display(ReportText('This subsection provides objectives and assumptions. '*100))

sub_methodology = Subsection('Methodology')
CurrentContainer(sub_methodology)
display(ReportText('This subsection provides methodology / model description / parameters. '*100))"""

math_str = """
from sympy import Eq, Symbol, symbols
from dynpy.utilities.report import *


sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

sec_description = Subsection('Description of dynamic model')
CurrentContainer(sec_description)
display(ReportText('This subsection provides description of the model. '*10))

from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys

dyn_sys = DynamicSys()
display(dyn_sys.as_picture())

display(ReportText('Summary of the subsection highlighting its major achievements. '*10))

sub_model_code = Subsection('Virtual model of the object')
CurrentContainer(sub_model_code)

from dynpy.utilities.report import ObjectCode

display(ReportText('This subsection provides code of the model. '*10))
display(ObjectCode(DynamicSys))
display(ReportText('Summary of the subsection highlighting its major achievements. '*10))


sec_lagrangian = Subsection('Lagrangian and derivatives')
CurrentContainer(sec_lagrangian)
display(ReportText('This subsection provides calculation of lagrangian and derivatives. '*10))

lagrangian = dyn_sys.L[0]
lagrangian

t=dyn_sys.ivar


for coord in dyn_sys.q:

    display(ReportText(f'Calculations of derivatives for ${vlatex(coord)}$'))


    vel=coord.diff(t)
    diff1 = lagrangian.diff(vel)
    diff2 = diff1.diff(t)
    diff3 = lagrangian.diff(coord)


    diff1_sym = Symbol(f'\\\\frac{{\\\\partial L}}{{\\\\partial {vlatex(vel)}}}')
    diff2_sym = Symbol(f' \\\\frac{{d }}{{dt}}  {diff1_sym}'  )
    diff3_sym = Symbol(f'\\\\frac{{\\\\partial L}}{{\\\\partial {vlatex(coord)}}}')


    display(SympyFormula(  Eq(diff1_sym,diff1)  ))
    display(SympyFormula(  Eq( diff2_sym ,diff2)  ))
    display(SympyFormula(  Eq( diff3_sym ,  diff3)  ))

display(ReportText('Outcomes of governing equations analysis. '*10))


sec_equations = Subsection('Equations of motion')
CurrentContainer(sec_equations)


display(ReportText('This subsection provides calculation of equations of motion and solution of the system. '*10))

ds1=dyn_sys.eoms

for eq1 in ds1.as_eq_list():
    display(SympyFormula(eq1.simplify()))

ds2=dyn_sys.linearized()._ode_system.general_solution

for eq2 in ds2.as_eq_list():
    display(SympyFormula(eq2.simplify()))

ds3=dyn_sys.linearized()._ode_system.steady_solution

for eq3 in ds3.as_eq_list():
    display(SympyFormula(eq3.simplify()))

ds4=dyn_sys.linearized().eoms.solution
ds4_eqns = ds4.as_eq_list()

for eq4 in ds4_eqns:
    display(SympyFormula(eq4.simplify()))


display(ReportText(f'Outcomes of governing equations {AutoMarker(ds4_eqns[0])} {AutoMarker(ds4_eqns[1])}) analysis.'))

sec_math_desc = Subsection('Section that contains all symbols descriptions')
CurrentContainer(sec_math_desc)



E_K,F = symbols('E_K,F')
descriptions = {
E_K:r"Kinetic energy",
F:r"Force",
}
syms_dict = descriptions
DescriptionsRegistry().set_descriptions({**syms_dict})
DescriptionsRegistry().reset_registry()
SymbolsDescription.set_default_header('where:')

display(ReportText('Symbols description is as follows: '))


display(SymbolsDescription(expr=E_K*F))"""

picture_str = """
sec_picture = Section('Section that presents pictures reporting')
CurrentContainer(sec_picture)

display(Picture('./dynpy/models/images/taipei101.png',caption = 'Caption of picture',width=global_width,height=global_height))"""

simulaion_str = """
from dynpy.utilities.report import *
from sympy import *


sec_simulation = Section('Section that contains simulation')
CurrentContainer(sec_simulation)

# Basics of ODESystem based simulations are covered in guide to ODESystem, use the followin call
# from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
# BasicsOfODESystemGuide()

sec_ODESystem = Subsection('ODESystem simulation')
CurrentContainer(sec_ODESystem)

display(ReportText('Firstly create an ODESystem or import it from dynpy.modes.odes.linear.py'))
display(ObjectCode('spring = SpringMassEps.from_reference_data()''))

display(ReportText('Secondly solve the equation using solution method:'))
display(ObjectCode('spring_sol = spring.solution'))

display(ReportText('Last step is data substitution and using numerized method. After that use $compute_solution$ to finalize the simulation:'))
spring_sol.subs(data_spring).with_ics([10,0]).numerized().compute_solution(t_span).plot()

# MSM_sim = Subsection('MSM simulation')
# CurrentContainer(MSM_sim)

# display(ReportText('Firstly create an MultiTimeScaleSolution system or import it from dynpy.modes.odes.linear.py'))
# display(ObjectCode('spring = SpringMassMSM.from_reference_data()'))

# display(ReportText('Secondly solve the equation using solution method:'))
# display(ObjectCode('spring_sol = spring.solution'))

# display(ReportText('Last step is data substitution and using numerized method. After that use $compute_solution$ to finalize the simulation:'))
# spring_sol.subs(data_spring).with_ics([10,0]).numerized().compute_solution(t_span).plot()"""

veryfication_str = """
from sympy import *
from dynpy.utilities.adaptable import *
from dynpy.utilities.report import CurrentContainer, ReportText
from sympy import *
from sympy.physics import units




from dynpy.utilities.adaptable import *
from dynpy.utilities.report import CurrentContainer, ReportText
import pandas as pd

t = Symbol('t')
power = Symbol('P')
work = Symbol('W')

unit_dict = {
t: units.second,
power: units.watt,
work: units.joule,
}


LatexDataFrame.set_default_units(unit_dict)



sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept. '*200))



### SELECT ONE CASE
## 1st CASE
## import of external data if path "./data_from_ml/test_1.csv" exists

#test_data = './data_from_ml/test_1.csv'
#df = pd.read_csv(filename, header = None)

## 2st CASE
df = pd.DataFrame({t:[0,1,2],'a':[2,3,4],'b':[8,7,6]}).set_index(t)

df_cols = df.set_axis([power,work],axis=1)

data_tab = TimeDataFrame(df_cols).to_latex_dataframe().reported(caption='Caption')
graph = TimeDataFrame(df_cols).to_pylatex_tikz().in_figure(caption='Caption')



display(data_tab)

display(ReportText('Obtained data analysis. '*200))

display(graph)

display(ReportText('Description of verifications outcomes. '*200))"""

conclusion_str = """
sec_conclusion = Section('Section that contains final conclusions')
CurrentContainer(sec_conclusion)

display(ReportText('Conclusions '*200))"""

symbols_description_str = """
sec_symbols = Section('Section that contains all symbols descriptions')
CurrentContainer(sec_symbols)

E_K,F = symbols('E_K,F')
descriptions = {
E_K:r"Kinetic energy",
F:r"Force",
}
syms_dict = descriptions
DescriptionsRegistry().set_descriptions({**syms_dict})
DescriptionsRegistry().reset_registry()
SymbolsDescription.set_default_header('  ')

display(SymbolsDescription({**syms_dict}))"""

document_str = """
# Creating file
# Be sure *output* folder is in the current directory

thesis_name = './output/report_name' #path for report file

doc = WutThesis(thesis_name)
# Bibliography of quotations
### Select one bibligraphy managment system
####### BibLatex

doc.preamble.append(Package('biblatex',arguments=["backend=biber","sorting=none"]))
doc.preamble.append(Command('addbibresource','elementy_bibliagrafia.bib'))
####### Natbib
#doc.preamble.append(Package('natbib')
#doc.preamble.append(Command('bibliographystyle','unsrt'))

## Units
doc.preamble.append(Package('siunitx'))

# TOC

doc.append(Command('includepdf{./Images/Front_Page.pdf}')) #includes front page##path shouldn't be changed, it must look like ./Images/Your_file.pdf
doc.append(Command('cleardoublepage'))
doc.append(Command('includepdf{./Images/oswiadczenie_autora_pracy.pdf}'))#path shouldn't be changed, it must look like ./Images/Your_file.pdf
doc.append(Command('pagestyle{plain}'))
doc.append(Command('cleardoublepage'))
doc.append(Command('includepdf{./output/oswiadczenie_biblioteczne.pdf}'))#path shouldn't be changed, it must look like ./Images/Your_file.pdf
doc.append(Command('pagestyle{plain}'))
doc.append(Command('cleardoublepage'))

doc.append(Command('tableofcontents')) #adds TOC

doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)

doc.append(sub_model_code)
doc.append(sec_math_desc)
doc.append(sec_simulation)

doc.append(sec_picture)
doc.append(sec_verification)
doc.append(sec_tables)
doc.append(sec_symbols)
doc.append(sec_conclusion)


### BibLatex
#doc.append(Command('printbibliography',arguments=["title={Bibliography}"])) - argument is to improve
doc.append(Command('printbibliography',options=[NoEscape("title={Bibliography}")]))

### Natbib
#doc.append(Command('bibliography',arguments=["references"])) # .bib file as "references"


#FIGURES LIST

doc.append(Command('addcontentsline{toc}{section}{List of figures}'))
doc.append(Command('listoffigures'))
doc.append(Command('pagestyle{plain}'))
doc.append(Command('newpage'))

#TABLES LIST

doc.append(Command('addcontentsline{toc}{section}{List of tables}'))
doc.append(Command('renewcommand{\listtablename}{List of tables}'))
doc.append(Command('listoftables'))
doc.append(Command('pagestyle{plain}'))
doc.append(Command('newpage'))

# Generating file
doc.generate_pdf(clean_tex=True)"""


class ReportMethods:

    _reported_object = None

    @classmethod
    def base_setup(cls):

        preliminary_str = """

#Examplary setup is as follows:

##
## Imports



    from dynpy.utilities.report import *
    from dynpy.utilities.templates.document import Guide

    doc = Guide('./output/report_name')


## CELL 2
## Text reporting

    #!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
    #!!!       BECAUSE OF NEEDED IMPORTS    !!!#

    sec_text = Section('Section that presents text reporting')
    CurrentContainer(sec_text)

    display(ReportText('Exemplary text'*100))

    #will not work in some projects, restrict usage
    display(Markdown('Formatted text'*100))


## CELL 3
## Math

    #!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
    #!!!       BECAUSE OF NEEDED IMPORTS    !!!#

    from sympy import Eq, Symbol, symbols

    sec_formula = Section('Section that presents formulas reporting')
    CurrentContainer(sec_formula)

    display(ReportText('Mathematical formulas are reported with the support of sympy and it\\'s symbols.'))

    a,b = symbols('a b')
    display(SympyFormula(Eq(a,b)))


## CELL 4
## Picture

    #!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
    #!!!       BECAUSE OF NEEDED IMPORTS    !!!#

    sec_picture = Section('Section that presents pictures reporting')
    CurrentContainer(sec_picture)

    display(Picture('./dynpy/models/images/taipei101.png',caption = 'Caption of picture'))



## CELL 5
## Document

    #!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
    #!!!       BECAUSE OF NEEDED IMPORTS    !!!#

    # Creating file
    # Be sure *output* folder is in the current directory

    guide_name = './output/report_name' #path for report file

    doc = Guide(guide_name)
    doc.append(sec_text) # adding certain sections
    doc.append(sec_formula)
    doc.append(sec_picture)
    # Generating file
    doc.generate_pdf(clean_tex=True)



"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = """
#Code below is the exact same, as the one presented in the beginning of the guide. Duplicated for easier copying and pasting.

#Method presented below initializes a simple document and prepares it for content addition.

#Good practice here is to allocate 1 section per 1 cell

## ############### CELL 1 ###########################
## Imports

from dynpy.utilities.report import *
from dynpy.utilities.templates.document import Guide

doc = Guide('./output/report_name')


## ############### CELL 2 ###########################
## Text reporting

sec_text = Section('Section that presents text reporting')
CurrentContainer(sec_text)

display(ReportText('Exemplary text'*100))

#will not work in some projects, restrict usage
display(Markdown('Formatted text'*100))


## ############### CELL 3 ###########################
## Math

from sympy import Eq, Symbol, symbols

sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

display(ReportText('Mathematical formulas are reported with the support of sympy and it\\'s symbols.'))

a,b = symbols('a b')
display(SympyFormula(Eq(a,b)))

## ############### CELL 4 ###########################
## Picture

sec_picture = Section('Section that presents pictures reporting')
CurrentContainer(sec_picture)

display(Picture('./dynpy/models/images/taipei101.png',caption = 'Caption of picture'))



## ############### CELL 5 ###########################
## Document

# Creating file
# Be sure *output* folder is in the current directory

guide_name = './output/report_name' #path for report file

doc = Guide(guide_name)
doc.append(sec_text) # adding certain sections
doc.append(sec_formula)
doc.append(sec_picture)
# Generating file
doc.generate_pdf(clean_tex=True)

"""
        return ObjectCode(preliminary_str)

    @property
    def _report_components(self):

        from ..components.guides import en as guide_comp

        comp_list = [
            # mech_comp.TitlePageComponent,
            # mech_comp.SchemeComponent,
            # mech_comp.ExemplaryPictureComponent,
            # mech_comp.KineticEnergyComponent,
            # mech_comp.PotentialEnergyComponent,
            # mech_comp.LagrangianComponent,
            # mech_comp.GoverningEquationComponent,
            # #mech_comp.FundamentalMatrixComponent,
            # mech_comp.GeneralSolutionComponent,
            # #mech_comp.SteadySolutionComponent,
        ]

        return comp_list

    @property
    def default_reported_object(self):

        return None

    @property
    def reported_object(self):

        reported_obj = self._reported_object

        if reported_obj is None:
            return self.default_reported_object
        else:
            return reported_obj

    @reported_object.setter
    def reported_object(self, value):

        self._reported_object = value

    def append_components(self, reported_object=None):

        # self.reported_object = reported_object

        doc = self

        for comp in self._report_components:
            doc.append(comp(self.reported_object))

        return None


class Guide(Document, ReportMethods):
    """
    A class to generate a structured LaTeX document with pre-configured packages and settings.

    Inherits from `Document` and provides additional methods to simplify LaTeX report creation.

    Attributes:
        _documentclass (str): The LaTeX document class (default is 'article').
        latex_name (str): Name of the document.
        packages (List[Union[Package, Command]]): List of LaTeX packages and commands to include in the document.
        _reported_object (Optional[object]): The object being reported on, if any.

    Exemplary Usage:
        >>> from dynpy.utilities.report import *
        >>> from dynpy.utilities.documents.guides import Guide

        >>> doc = Guide('./output/sample_report', title="Sample Report")

        >>> section = Section('Exemplary section name')
        >>> CurrentContainer(section)

        >>> display(Markdown(''' Exemplary Markdown text in the section '''))
        >>> display(ReportText(' Exemplary text appended into section '))

        >>> doc.append(section)

        >>> doc.generate_pdf(clean_tex=True)



    Example of Customization of the document properties:
        >>> from dynpy.utilities.report import *
        >>> from dynpy.utilities.documents.guides import Guide

        >>> custom_geometry = ['lmargin=20mm', 'rmargin=20mm', 'top=25mm', 'bmargin=25mm']

        >>> doc = Guide(
        >>>     default_filepath='./output/custom_report',
        >>>     title='Custom Report',
        >>>     geometry_options=custom_geometry
        >>> )

        >>> section = Section('Exemplary Custom Section Name')

        >>> doc.append(section)

        >>> doc.generate_pdf(clean_tex=True)
    """

    _documentclass = "article"
    latex_name = "document"
    packages = [
        Package(
            "geometry",
            options=[
                "lmargin=25mm",
                "rmargin=25mm",
                "top=30mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        #Package("microtype"),
        Package("authoraftertitle"),
        Package("polski", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Command("pagestyle", arguments=["fancy"]),
        Command("fancyhf", arguments=[""]),
        Command("fancyhead", arguments=["DynPy Team"], options=["R"]),
        Command("fancyhead", arguments=["Mechanical vibration, 2023"], options=["L"]),
        Command("fancyfoot", arguments=[NoEscape("\\thepage")], options=["C"]),
    ]

    def __init__(
        self,
        default_filepath: str = "default_filepath",
        title: str = "Basic title",
        reported_object: Optional[object] = None,
        *,
        documentclass: Optional[str] = None,
        document_options: Optional[List[str]] = None,
        fontenc: str = "T1",
        inputenc: str = "utf8",
        font_size: str = "normalsize",
        lmodern: bool = False,
        textcomp: bool = True,
        microtype: bool = True,
        page_numbers: bool = True,
        indent: Optional[Union[str, int]] = None,
        geometry_options: Optional[List[str]] = None,
        data: Optional[dict] = None,
    ):
        """
        Initialize the Guide class with optional customization.

        Args:
            default_filepath (str): Path for the generated document.
            title (str): Title of the document.
            reported_object (Optional[object]): Object being reported on, if any.
            documentclass (Optional[str]): LaTeX document class (e.g., 'article').
            document_options (Optional[List[str]]): Options for the document class.
            fontenc (str): Font encoding (default 'T1').
            inputenc (str): Input encoding (default 'utf8').
            font_size (str): Font size (default 'normalsize').
            lmodern (bool): Whether to use Latin Modern fonts.
            textcomp (bool): Whether to use the `textcomp` package.
            microtype (bool): Whether to use the `microtype` package.
            page_numbers (bool): Whether to include page numbers.
            indent (Optional[Union[str, int]]): Indentation settings.
            geometry_options (Optional[List[str]]): Geometry options for page layout.
            data (Optional[dict]): Additional data for customization.
        """

        if documentclass is not None:
            self._documentclass

        self._reported_object = reported_object

        super().__init__(
            default_filepath=default_filepath,
            documentclass=self._documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )

        #         label=self.label
        self.title = "Mechanical vibration"
        # self.packages.append(Command('title', arguments=[NoEscape(self.title)]))
        # self.packages.append(Command('author', arguments=['DynPy Team']))
        self.packages.append(Command("date", arguments=[NoEscape("\\today")]))
        # self.append(Command('maketitle'))
        self.append(NewPage())
        # tu implementować co tam potrzeba
        self.append_components()


class CaseTemplate(Document, ReportMethods):

    latex_name = "document"
    _documentclass = "report"

    packages = [
        Package("microtype"),
        Package("polski", options=["MeX"]),
        Package(
            "geometry",
            options=[
                "lmargin=25mm",
                "rmargin=25mm",
                "top=30mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Command("pagestyle", arguments=["fancy"]),
        Command("fancyhf", arguments=[""]),
        Command("fancyhead", arguments=["DynPy Team"], options=["R"]),
        Command("fancyhead", arguments=["Mechanical Vibration, 2021"], options=["L"]),
        Command("fancyfoot", arguments=[NoEscape("\\thepage")], options=["C"]),
        # Command('newcommand{\praca}', arguments=['Praca dyplomowa']),
        # Command('newcommand{\dyplom}', arguments=['Inżynierska']),
        # Command('newcommand{\kierunek}', arguments=['Wpisać kierunek']),
        # Command('newcommand{\specjalnosc}', arguments=['Wpisać specjalność']),
        # Command('newcommand{\\autor}', arguments=['Imię i nazwisko autora']),
        # Command('newcommand{\opiekun}', arguments=['Wpisać opiekuna']),
        # Command('newcommand{\promotor}', arguments=['Wpisać promotora']),
        # Command('newcommand{\konsultant}', arguments=['Wpisać konsultanta']),
        # Command('newcommand{\\tytul}', arguments=['Wpisać tytuł pracy dyplomowej po polsku']),
        # Command('newcommand{\\title}', arguments=['Wpisać tytuł pracy dyplomowej po angielsku']),
        # Command('newcommand{\supervisor}', arguments=['dr inż. Bogumił Chiliński']),
        # Command('newcommand{\\rok}', arguments=['Rok składania pracy']),
        # Command('newcommand{\kluczowe}', arguments=['Słowa kluczowe: Wpisać słowa kluczowe po polsku']),
        # Command('newcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']),
    ]

    # \renewcommand{\headrulewidth}{1pt}
    # \renewcommand{\footrulewidth}{1pt}

    def __init__(
        self,
        default_filepath="default_filepath",
        reported_object=None,
        *,
        documentclass=None,
        document_options=None,
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,  # ['lmargin=25mm', 'rmargin=25mm',  'top=20mm', 'bmargin=25mm', 'headheight=50mm'],
        data=None,
    ):

        if documentclass is not None:
            self._documentclass

        self._reported_object = reported_object

        super().__init__(
            default_filepath=default_filepath,
            documentclass=self._documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )

        self.append_components()


class ExampleTemplate(Guide):
    pass


class BPASTSPaper(Document):

    latex_name = "document"
    packages = [
        #                     Package('natbib', options=['numbers']),
        Package("booktabs"),
        Package("float"),
        Package("standalone"),
        Package("siunitx"),
        #                     Package('bpasts', options=['accepted']),
        Package("bpasts"),
        Package("t1enc"),
        Package("amsmath"),
        Package("amssymb"),
        Package("amsfonts"),
        Package("graphicx"),
        Package("flushend"),
        Package("textcomp"),
        Package("xcolor"),
        Package(
            "hyperref",
            ["colorlinks=true", "allcolors=bpastsblue", NoEscape("pdfborder={0 0 0}")],
        ),
    ]

    abtitle = "Paper for BPASTS"
    abauthor = "Authors"
    title = "Basic title"

    def __init__(
        self,
        default_filepath="default_filepath",
        title=None,
        *,
        documentclass="article",
        document_options=["10pt", "twoside", "twocolumn", "a4paper"],  # for submission
        fontenc=None,
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=False,
        microtype=False,
        page_numbers=True,
        indent=None,
        geometry_options=[
            "inner=30mm",
            "outer=20mm",
            "bindingoffset=10mm",
            "top=25mm",
            "bottom=25mm",
        ],  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        data=None,
    ):

        if title is not None:
            self.title = title

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )

        self.preamble.append(Command("abtitle", arguments=[self.abtitle]))
        self.preamble.append(NoEscape("\\title{" + f"{self.title}" + "}"))
        self.preamble.append(NoEscape("\\abauthor{" + f"{self.abauthor}" + "}"))

        self.preamble.append(NoEscape("%%%% EDITORS SECTION"))
        self.preamble.append(NoEscape("\\vol{XX} \\no{Y} \\year{2024}"))
        self.preamble.append(NoEscape("\\setcounter{page}{1}"))
        self.preamble.append(NoEscape("\\doi{10.24425/bpasts.yyyy.xxxxxx}"))
        self.preamble.append(NoEscape("%%%%%%%%%%%%%%%%%%%%"))
        self.append(Command("maketitle"))
        # self.append(NewPage())
        # tu implementować co tam potrzeba

    def cwd_setup(self):
        cwd = os.getcwd()

        source_path = (
            f"/home/user/Shared files/modules/dynpy/utilities/documents/Definitions"
        )
        path1 = f"{cwd}/bpast.sty"
        path2 = f"{cwd}/output/bpast.sty"

        if os.path.exists(path1) == False:
            shutil.copytree(source_path, path1)
        if os.path.exists(path2) == False:
            shutil.copytree(source_path, path2)

    def authors(self, nameno, affiliation=None, coremail=None, corno=0):

        # nameno - dict; of author name (string) and affiliation number (0 to x integer)
        # affiliation - dcit; affiliation number (0 to x integer) and affiliation (string)
        # coremail - str; corresponding author email
        # corno - int; which of the authors in the dict is the corresponding author, count starting from 0

        # HOW TO USE?

        # Doc1 = BPASTSPaper(default_filepath='./output/method_test',title=NoEscape('''Your title'''))
        # author_names={'Damian Sierociński':1,'Bogumił Chiliński':1, 'Karolina Ziąbrowska':2, 'Jakub Szydłowski':2}
        # author_affiliations={1:'Institute of Machine Design Fundamentals, Faculty of Automotive and Construction Machinery Engineering, Warsaw University of Technology',2:'Faculty of Automotive and Construction Machinery Engineering, Warsaw University of Technology'}
        # correspondence='damian.sierocinski@pw.edu.pl'
        # Doc1.authors(nameno=author_names,affiliation=author_affiliations,coremail=correspondence)

        counter = 0
        auth_string = ""
        addr_string = ""
        for name, no in nameno.items():
            auth_string = auth_string + name + "$^{" + f"{no}" + "}$"
            if counter == corno:
                auth_string = auth_string + "\email{" + coremail + "}"
            counter = counter + 1

            if counter != len(nameno):
                auth_string = auth_string + ", "

        counter_addr = 1
        for no, addr in affiliation.items():
            addr_string = addr_string + "$^" + f"{no}" + "$" + addr
            if counter_addr != len(affiliation):
                addr_string = addr_string + ", "
            counter_addr = counter_addr + 1

        self.preamble.append(Command("author", arguments=[NoEscape(auth_string)]))
        self.preamble.append(Command("Address", arguments=[NoEscape(addr_string)]))

    def abstract(self, container=None):
        for num, val in enumerate(container):
            self.preamble.append(Command("Abstract", arguments=[ReportText(val)]))

    def keywords(self, keywords=None):
        self.preamble.append(Command("Keywords", arguments=[keywords]))

    @classmethod
    def base_setup(cls, create_file: bool = False):
        """
        Sets up the base template for the BPASTSPaper document.

        This method provides an easy way to initialize a thesis document with the necessary sections, structure, and configuration.

        Args:
            create_file (bool): If True, creates a base setup environment including necessary directories and files.

        Returns:
            Optional[str]: Returns setup content if `create_file` is False, otherwise creates the environment.

        Usage Example:
            >>> from dynpy.utilities.documents.document import BPASTSPaper
            >>> BPASTSPaper.base_setup()

        When to use:
            - Use this method when starting a new BPASTSPaper document.
            - If `create_file=True`, it will generate all required directories and files for a structured thesis.
            - If `create_file=False`, it returns a string with setup instructions that can be used manually.
        """

        if create_file is True:
            return cls._create_base_setup_env()

        # doc = WutThesis('./output/thesis_name')
        # doc.preamble.append(NoEscape(r'\\usepackage[MeX]{polski}')) #to set polish as main language"""

        thesis_introduction_str = """#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_intro = Section('Section that presents text reporting')
CurrentContainer(sec_intro)

sub_problem_outline = Subsection('Outline of the problem')
CurrentContainer(sub_problem_outline)
display(ReportText('This subsection provides information about investigated problem. '*100))

sub_obj_assum = Subsection('Objectives and assumptions')
CurrentContainer(sub_obj_assum)
display(ReportText('This subsection provides objectives and assumptions. '*100))

sub_SOT = Subsection('State of the art')
CurrentContainer(sub_SOT)
display(ReportText('This subsection provides state of the art \\cite{pandas}. '*100))

sub_methodology = Subsection('Methodology')
CurrentContainer(sub_methodology)
display(ReportText('This subsection provides methodology. '*100))"""

        math_str = '''from sympy import Eq, Symbol, symbols
from dynpy.utilities.report import *


sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

sec_description = Subsection('Description of dynamic model')
CurrentContainer(sec_description)
display(ReportText('This subsection provides description of the model. '*10))

from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys

dyn_sys = DynamicSys()
display(dyn_sys.as_picture())

display(ReportText('Summary of the subsection highlighting its major achievements. '*10))

sub_model_code = Subsection('Virtual model of the object')
CurrentContainer(sub_model_code)

from dynpy.utilities.report import ObjectCode

display(ReportText('This subsection provides code of the model. '*10))
display(ObjectCode(DynamicSys))
display(ReportText('Summary of the subsection highlighting its major achievements. '*10))


sec_lagrangian = Subsection('Lagrangian and derivatives')
CurrentContainer(sec_lagrangian)
display(ReportText('This subsection provides calculation of lagrangian and derivatives. '*10))

lagrangian = dyn_sys.L[0]
lagrangian

t=dyn_sys.ivar


for coord in dyn_sys.q:

    display(ReportText(f'Calculations of derivatives for ${vlatex(coord)}$'))


    vel=coord.diff(t)
    diff1 = lagrangian.diff(vel)
    diff2 = diff1.diff(t)
    diff3 = lagrangian.diff(coord)


    diff1_sym = Symbol(f'\\\\frac{{\\\\partial L}}{{\\\\partial {vlatex(vel)}}}')
    diff2_sym = Symbol(f' \\\\frac{{d }}{{dt}}  {diff1_sym}'  )
    diff3_sym = Symbol(f'\\\\frac{{\\\\partial L}}{{\\\\partial {vlatex(coord)}}}')


    display(SympyFormula(  Eq(diff1_sym,diff1)  ))
    display(SympyFormula(  Eq( diff2_sym ,diff2)  ))
    display(SympyFormula(  Eq( diff3_sym ,  diff3)  ))

display(ReportText('Outcomes of governing equations analysis. '*10))


sec_equations = Subsection('Equations of motion')
CurrentContainer(sec_equations)


display(ReportText('This subsection provides calculation of equations of motion and solution of the system. '*10))

ds1=dyn_sys.eoms
ds1_eqns=ds1.as_eq_list()

for eq1 in ds1.as_eq_list():
    display(SympyFormula(eq1.simplify()))

if ds1.is_solvable():

    ds2=dyn_sys.linearized()._ode_system.general_solution

    for eq2 in ds2.as_eq_list():
        display(SympyFormula(eq2.simplify()))

    ds3=dyn_sys.linearized()._ode_system.steady_solution

    for eq3 in ds3.as_eq_list():
        display(SympyFormula(eq3.simplify()))

    ds4=dyn_sys.linearized().eoms.solution
    ds4_eqns = ds4.as_eq_list()

    for eq4 in ds4_eqns:
        display(SympyFormula(eq4.simplify()))


display(ReportText(f'Outcomes of governing equations {AutoMarker(ds1_eqns[0])} {AutoMarker(ds1_eqns[-1])}) analysis.'))

sec_math_desc = Subsection('Section that contains all symbols descriptions')
CurrentContainer(sec_math_desc)



E_K,F = symbols('E_K,F')
descriptions = {
E_K:r"Kinetic energy",
F:r"Force",
}
syms_dict = descriptions
DescriptionsRegistry().set_descriptions({**syms_dict})
DescriptionsRegistry().reset_registry()
SymbolsDescription.set_default_header('where:')

display(ReportText('Symbols description is as follows: '))


display(SymbolsDescription(expr=E_K*F))


display(ObjectCode("""
# Guide
from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
BasicsOfODESystemGuide();

"""))


# Basics of DynSys usage based simulations are covered in guide to DynSys usage, use the followin call
# from dynpy.utilities.documents.guides import UsageOfDynamicSystemsGuide
# UsageOfDynamicSystemsGuide()
'''

        picture_str = """sec_picture = Section('Section that presents pictures reporting')
CurrentContainer(sec_picture)

display(Picture('./dynpy/models/images/taipei101.png',caption = 'Caption of picture'))"""

        simulaion_str = """from dynpy.utilities.report import *
from sympy import *


sec_simulation = Section('Section that contains simulation')
CurrentContainer(sec_simulation)

# Basics of ODESystem based simulations are covered in guide to ODESystem, use the followin call
# from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
# BasicsOfODESystemGuide()

sec_ODESystem = Subsection('ODESystem simulation')
CurrentContainer(sec_ODESystem)

display(ReportText('Firstly create an ODESystem or import it from dynpy.modes.odes.linear.py'))
display(ObjectCode('spring = SpringMassEps.from_reference_data()'))

display(ReportText('Secondly solve the equation using solution method:'))
display(ObjectCode('spring_sol = spring.solution'))

display(ReportText('Last step is data substitution and using numerized method. After that use $compute_solution$ to finalize the simulation:'))
data_spring = dyn_sys.get_numerical_parameters()
#ds1.subs(data_spring).with_ics([10,0]).numerized().compute_solution(t_span).plot()

# MSM_sim = Subsection('MSM simulation')
# CurrentContainer(MSM_sim)

# display(ReportText('Firstly create an MultiTimeScaleSolution system or import it from dynpy.modes.odes.linear.py'))
# display(ObjectCode('spring = SpringMassMSM.from_reference_data()'))

# display(ReportText('Secondly solve the equation using solution method:'))
# display(ObjectCode('spring_sol = spring.solution'))

# display(ReportText('Last step is data substitution and using numerized method. After that use $compute_solution$ to finalize the simulation:'))
# spring_sol.subs(data_spring).with_ics([10,0]).numerized().compute_solution(t_span).plot()"""

        veryfication_str = """from sympy import *

from sympy import *
from sympy.physics import units




from dynpy.utilities.adaptable import *
from dynpy.utilities.report import CurrentContainer, ReportText
import pandas as pd

t = Symbol('t')
power = Symbol('P')
work = Symbol('W')

unit_dict = {
t: units.second,
power: units.watt,
work: units.joule,
}

LatexDataFrame.set_default_units(unit_dict)



sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept. '*200))



### SELECT ONE CASE
## 1st CASE
## import of external data if path "./data_from_ml/test_1.csv" exists

#test_data = './data_from_ml/test_1.csv'
#df = pd.read_csv(test_data, header = None)

#data_tab = TimeDataFrame(df).to_latex_dataframe().reported(caption='Caption')
#graph=data_tab.iloc[:,[1]]

#display(data_tab)
display(ReportText('Obtained data analysis. '*200))
#display(graph)

## 2st CASE
## Creating graph from a
df = pd.DataFrame({t:[0,1,2],'a':[2,3,4],'b':[8,7,6]}).set_index(t)

df_cols = df.set_axis([power,work],axis=1)

data_tab = TimeDataFrame(df_cols).to_latex_dataframe().reported(caption='Caption')
graph = TimeDataFrame(df_cols).to_pylatex_tikz().in_figure(caption='Caption')

display(data_tab)
display(ReportText('Obtained data analysis. '*200))

display(graph)

## 3rd CASE
## dictionary data presentation in DataFrame
#dictionary ={
#     power:300000,
#     work:5000
# }
#display(
#     LatexDataFrame.formatted(
#         data=dictionary, #import data from a dictionary
#         index=['Value'] #set rows title
#     ).map(lambda x: f'${latex(x)}$')
#     .rename_axis('Parameter', axis=1) #set name for a column
#     .transpose()  # This swaps rows and columns, making the DataFrame vertical
#     .reported(caption='A caption for your data frame')
# )
#display(ReportText('Obtained data analysis. '*200))


## 4th CASE
## creating a graph based on simulation
#from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys
#F=DynamicSys.F
#g=DynamicSys.g
#m=DynamicSys.m
#k=DynamicSys.k
#slownik ={
#F:2,
#g:10,
#m:5,
#k:100,
#}
#t_span = np.linspace(0.0,1,30)
#sym=DynamicSys().eoms.subs(slownik)
#wynik= sym.numerized(backend='numpy').compute_solution(t_span,[0.0,0.0])#.plot()
#wynik_tab = TimeDataFrame(wynik).to_latex_dataframe().reported(caption='Caption')
#display(wynik_tab)
#display(ReportText('Obtained data analysis. '*200))
#wynik.plot()


display(ReportText('Description of verifications outcomes. '*200))"""

        conclusion_str = """sec_conclusion = Section('Section that contains final conclusions')
CurrentContainer(sec_conclusion)

display(ReportText('Conclusions '*200))"""

        symbols_description_str = """sec_symbols = Section('Section that contains all symbols descriptions')
CurrentContainer(sec_symbols)

E_K,F = symbols('E_K,F')
descriptions = {
E_K:r"Kinetic energy",
F:r"Force",
}
syms_dict = descriptions
DescriptionsRegistry().set_descriptions({**syms_dict})
DescriptionsRegistry().reset_registry()
SymbolsDescription.set_default_header('  ')

display(SymbolsDescription({**syms_dict}))"""

        document_str = """# Creating file
# Be sure *output* folder is in the current directory

paper_name = './output/report_name' #path for report file
global_width=('8cm')#setting picture default size
global_height=('6cm')
doc = BPASTSPaper(paper_name)
# Bibliography of quotations
### Select one bibligraphy managment system
####### BibLatex

doc.preamble.append(Package('biblatex',["backend=biber","sorting=none"]))
doc.preamble.append(Command('addbibresource','elementy_bibliagrafia.bib'))
####### Natbib
#doc.preamble.append(Package('natbib')
#doc.preamble.append(Command('bibliographystyle','unsrt'))

## Units
doc.preamble.append(Package('siunitx'))




doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)

doc.append(sub_model_code)
doc.append(sec_math_desc)
doc.append(sec_simulation)

doc.append(sec_picture)
doc.append(sec_verification)
#doc.append(sec_tables)
doc.append(sec_symbols)
doc.append(sec_conclusion)




### Natbib
doc.append(Command('bibliography',arguments=["references"])) # .bib file as "references"


#FIGURES LIST

doc.append(Command('addcontentsline{toc}{section}{List of figures}'))
doc.append(Command('listoffigures'))
doc.append(Command('pagestyle{plain}'))
doc.append(Command('newpage'))

#TABLES LIST

doc.append(Command('addcontentsline{toc}{section}{List of tables}'))
doc.append(Command('renewcommand{\listtablename}{List of tables}'))
doc.append(Command('listoftables'))
doc.append(Command('pagestyle{plain}'))
doc.append(Command('newpage'))

# Generating file
doc.generate_pdf(clean_tex=False)"""

        #         preliminary_str=(f"""
        # =======
        preliminary_str = f"""

Examplary setup is as follows:

#References to guide imports needed to include - to see the list of available guides insert and run the code below:
```python
from dynpy.utilities.creators import list_of_guides
list_of_guides()
```

## CELL 1
## Imports

```python
{imports_str}


```



## CELL 2
## Thesis introduction

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{thesis_introduction_str}
```



## CELL 3
## Math

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{math_str}
```


## CELL 4
## Picture


```python

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{picture_str}

```


## CELL 5
## Simulation

Basics of ODESystem based simulations are covered in guide to ODESystem, use the following call:

```python
from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
BasicsOfODESystemGuide();
```


```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#


{simulaion_str}

```


## CELL 6
## Veryfication

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{veryfication_str}

```

## CELL 7
## Conclusion

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{conclusion_str}

```

## CELL 9
## Symbols description

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{symbols_description_str}

```

## CELL 10
## Document

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{document_str}
```

"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = f"""
#Example:

#To prepare a simple document with text and images:

#Good practice here is to allocate 1 section per 1 cell

## CELL 1
## Imports

{imports_str}

## CELL 2
## Thesis introduction

{thesis_introduction_str}

## CELL 3
## Math

{math_str}

## CELL 4
## Picture

{picture_str}

## CELL 5
## Simulation

{simulaion_str}

## CELL 6
## Veryfication

{veryfication_str}

## CELL 7
## Conclusion

{conclusion_str}

## CELL 9
## Symbols description

{symbols_description_str}

## CELL 10
## Document

{document_str}

"""

        return ObjectCode(preliminary_str)


class APPAPaper(Document):

    latex_name = "document"
    packages = [
        #                     Package('natbib', options=['numbers']),
        Package("booktabs"),
        Package("float"),
        Package("standalone"),
        Package("siunitx"),
        #                     Package('bpasts', options=['accepted']),
        #Package("bpasts"),
        Package("t1enc"),
        Package("amsmath"),
        Package("amssymb"),
        Package("amsfonts"),
        Package("graphicx"),
        Package("flushend"),
        Package("textcomp"),
        Package("xcolor"),
#         Package(
#             "hyperref",
#             ["colorlinks=true", "allcolors=bpastsblue", NoEscape("pdfborder={0 0 0}")],
#         ),
    ]

    abtitle = "Paper for BPASTS"
    abauthor = "Authors"
    title = "Basic title"
    author = r"""
Author's Name Surname  - Co-author's Name Surname
"""
    abstract=r"""
The Abstract should not exceed 250 words. The Abstract should state the principal objectives and the scope of the investigation, as well as the methodology employed. It should summarize the results and state the principal conclusions. An effective abstract stands on its own — it can be understood fully even when made available without the full paper. To this end, avoid referring to figures or the bibliography in the abstract. Please introduce any acronyms the first time you use them in the abstract (if needed), and do so again in the full paper. About 4 to 6 significant key words should follow the abstract to aid indexing.
"""

    keywords = "key word, key word, key word, key word, key word"

    highlights=r"""
\item Highlights (4 to 6) are a short collection of bullet points that convey the core findings and provide readers with a quick textual overview of the article.
\item These four to six bullet points should describe the essence of the research (e.g. results or conclusions) and highlight what is distinctive about it.
\item See latest SV-JME papers for examples.
"""


    def __init__(
        self,
        default_filepath="default_filepath",
        title=None,
        *,
        documentclass="article",
        document_options=["10pt",
                          #"twoside", 
                        "twocolumn", "a4paper",
                         ],  # for submission
        fontenc=None,
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=False,
        microtype=False,
        page_numbers=True,
        indent=None,
        #geometry_options=[
        #    "inner=30mm",
        #    "outer=20mm",
        #    "bindingoffset=10mm",
        #    "top=25mm",
        #    "bottom=25mm",
        #],  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        geometry_options=None,
        data=None,
    ):

        if title is not None:
            self.title = title

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )



#         self.preamble.append(NoEscape(r"""
# % to be set by Editor:
# % ------------------------------------------------------------------------------------------------------------------
# %\renewcommand{\svCorr}{*Corr. Author's Address: Name of institution, Address, City, Country, \email{xxx.yyy@wwww.zzz}}
# \renewcommand{\svSV}{Strojniški vestnik - Journal of Mechanical Engineering 63(2017)3,  XXX-4}
# \renewcommand{\svAuthors}{Author-A, Author-B}
# \renewcommand{\svCyear}{2017}
# \renewcommand{\svDOI}{DOI: 10.5545/sv-jme.2017.4027}
# \svDates{2016-11-04}{2017-01-14}{2017-02-12}
# \renewcommand{\svType}{Review Paper}
# \setcounter{page}{1} % first page of the paper
# % ------------------------------------------------------------------------------------------------------------------
# """))


        self.preamble.append(Command("title", arguments=[self.title]))
        #self.preamble.append(NoEscape("\\title{" + f"{self.title}" + "}"))
        #self.preamble.append(NoEscape("\\abauthor{" + f"{self.abauthor}" + "}"))


        self.preamble.append(Command("author", arguments=[NoEscape(self.author)]))
        #self.preamble.append(Command('affil', 'Politechnika Warszawska'))

#         self.preamble.append(NoEscape(r"""

# %\date{\today\  -- \clock}
# \date{}

# \crop[cross,axes]
# %\crop[frame,axes]
# %\crop[off]"""))

        self.append(NoEscape(r"""
\twocolumn[
\maketitle
\begin{abstract}
"""))

        self.append(self.abstract)

        self.append(NoEscape(r"""
\end{abstract}]
"""))
#         self.append(Command("svKeywords", arguments=[NoEscape(self.keywords)]))

#         self.append(NoEscape(r"""
# \begin{svHigh}
# """))
#         self.append(NoEscape(self.highlights))

#         self.append(NoEscape(r"""

# \end{svHigh}

# \end{svHead}]
#"""))


        # self.append(NewPage())
        # tu implementować co tam potrzeba

    # def cwd_setup(self):
    #     cwd = os.getcwd()

    #     source_path = (
    #         f"/home/user/Shared files/modules/dynpy/utilities/documents/Definitions"
    #     )
    #     path1 = f"{cwd}/bpast.sty"
    #     path2 = f"{cwd}/output/bpast.sty"

    #     if os.path.exists(path1) == False:
    #         shutil.copytree(source_path, path1)
    #     if os.path.exists(path2) == False:
    #         shutil.copytree(source_path, path2)

class JoMEPaper(Document):

    latex_name = "document"
    packages = [
        #                     Package('natbib', options=['numbers']),
        Package("booktabs"),
        Package("float"),
        Package("standalone"),
        Package("siunitx"),
        #                     Package('bpasts', options=['accepted']),
        #Package("bpasts"),
        Package("t1enc"),
        Package("amsmath"),
        Package("amssymb"),
        Package("amsfonts"),
        Package("graphicx"),
        Package("flushend"),
        Package("textcomp"),
        Package("xcolor"),
#         Package(
#             "hyperref",
#             ["colorlinks=true", "allcolors=bpastsblue", NoEscape("pdfborder={0 0 0}")],
#         ),
    ]

    abtitle = "Paper for BPASTS"
    abauthor = "Authors"
    title = "Basic title"
    author = r"""
Author's Name Surname\svMark{1,*}  - Co-author's Name Surname\svMark{2} \\[1mm]
\svAffil{1}{Author's institution} \\
\svAffil{2}{Co-Author's institution}
"""
    abstract=r"""
The Abstract should not exceed 250 words. The Abstract should state the principal objectives and the scope of the investigation, as well as the methodology employed. It should summarize the results and state the principal conclusions. An effective abstract stands on its own — it can be understood fully even when made available without the full paper. To this end, avoid referring to figures or the bibliography in the abstract. Please introduce any acronyms the first time you use them in the abstract (if needed), and do so again in the full paper. About 4 to 6 significant key words should follow the abstract to aid indexing.
"""

    keywords = "key word, key word, key word, key word, key word"

    highlights=r"""
\item Highlights (4 to 6) are a short collection of bullet points that convey the core findings and provide readers with a quick textual overview of the article.
\item These four to six bullet points should describe the essence of the research (e.g. results or conclusions) and highlight what is distinctive about it.
\item See latest SV-JME papers for examples.
"""


    def __init__(
        self,
        default_filepath="default_filepath",
        title=None,
        *,
        documentclass="JoME",
        document_options=["10pt",
                          #"twoside", "twocolumn", "a4paper"
                         ],  # for submission
        fontenc=None,
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=False,
        microtype=False,
        page_numbers=True,
        indent=None,
        #geometry_options=[
        #    "inner=30mm",
        #    "outer=20mm",
        #    "bindingoffset=10mm",
        #    "top=25mm",
        #    "bottom=25mm",
        #],  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        geometry_options=None,
        data=None,
    ):

        if title is not None:
            self.title = title

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )



        self.preamble.append(NoEscape(r"""
% to be set by Editor:
% ------------------------------------------------------------------------------------------------------------------
%\renewcommand{\svCorr}{*Corr. Author's Address: Name of institution, Address, City, Country, \email{xxx.yyy@wwww.zzz}}
\renewcommand{\svSV}{Strojniški vestnik - Journal of Mechanical Engineering 63(2017)3,  XXX-4}
\renewcommand{\svAuthors}{Author-A, Author-B}
\renewcommand{\svCyear}{2017}
\renewcommand{\svDOI}{DOI: 10.5545/sv-jme.2017.4027}
\svDates{2016-11-04}{2017-01-14}{2017-02-12}
\renewcommand{\svType}{Review Paper}
\setcounter{page}{1} % first page of the paper
% ------------------------------------------------------------------------------------------------------------------
"""))


        self.preamble.append(Command("svTitle", arguments=[self.title]))
        self.preamble.append(NoEscape("\\title{" + f"{self.title}" + "}"))
        #self.preamble.append(NoEscape("\\abauthor{" + f"{self.abauthor}" + "}"))


        self.preamble.append(Command("author", arguments=[NoEscape(self.author)]))

        self.preamble.append(NoEscape(r"""

%\date{\today\  -- \clock}
\date{}

\crop[cross,axes]
%\crop[frame,axes]
%\crop[off]"""))

        self.append(NoEscape(r"""
\twocolumn[\begin{svHead}

\begin{svAbstract}
"""))

        self.append(self.abstract)

        self.append(NoEscape(r"""
\end{svAbstract}
"""))
        self.append(Command("svKeywords", arguments=[NoEscape(self.keywords)]))

        self.append(NoEscape(r"""
\begin{svHigh}
"""))
        self.append(NoEscape(self.highlights))

        self.append(NoEscape(r"""

\end{svHigh}

\end{svHead}]
"""))


        # self.append(NewPage())
        # tu implementować co tam potrzeba

    def cwd_setup(self):
        cwd = os.getcwd()

        source_path = (
            f"/home/user/Shared files/modules/dynpy/utilities/documents/Definitions"
        )
        path1 = f"{cwd}/bpast.sty"
        path2 = f"{cwd}/output/bpast.sty"

        if os.path.exists(path1) == False:
            shutil.copytree(source_path, path1)
        if os.path.exists(path2) == False:
            shutil.copytree(source_path, path2)



class Keywords(Environment):
    latex_name = "keyword"


class Abstract(Environment):
    latex_name = "abstract"


class ElsevierPaper(Document):
    """
    Doc = ElsevierTemplate(default_filepath='./output/paper')
    Doc.journal('Advances in Engineering Software')

    Doc.begin_frontmatter

    Doc.corresponding_author('Damian Sierociński','damian.sierocinski@pw.edu.pl')
    Doc.authors(['Bogumił Chiliński','Franciszek Gawiński','Piotr Przybyłowicz','Amadeusz Radomski'])
    Doc.abstract([NoEscape('\\textit{DynPy} is an open-source library implemented in \\textit{Python} programming language which aims to provide a versatile set of functionalities, that enable the user to model, solve, simulate, and report an analysis of a dynamic system in a single environment. In the paper examples for obtaining analytical and numerical solutions of the systems described with ordinary differential equations were presented. The assessment of solver accuracy was conducted utilising a model of a direct current motor and \\textit{MATLAB/Simulink} was used as a reference tool. The model was solved in \\textit{DynPy} with the hybrid analytical-numerical method and fully analytically, while in \\textit{MATLAB/Simulink} strictly numerical simulations were run. The comparison of the results obtained from both tools not only proved the credibility of the developed library but also showed its superiority in specific conditions. Moreover, the versatility of the tool was confirmed, as the whole process – from the definition of the model to the output of the formatted paper – was handled in the scope of a single \\textit{Jupyter} notebook.')])
    Doc.keywords(['engineering software','numerical simulations','analytical solution','electrical circuit',NoEscape('\\textit{Python} programming language')])
    Doc.title(NoEscape("An assessment of \\textit{DynPy} solver accuracy"))
    Doc.end_frontmatter
    """

    latex_name = "document"
    packages = [
        #                   Package('geometry',options=['lmargin=30mm', 'rmargin=30mm',  'top=25mm', 'bmargin=25mm', 'headheight=50mm']),
        #                   Package('microtype'),
        #                   Package('authoraftertitle'),
        #                   Package('listings'),
        #                   Package('titlesec'),
        #                   Package('fancyhdr'),
        #                   Package('graphicx'),
        #                   Package('indentfirst'),
        #                   Package('pdfpages'),
        Package("fontspec"),
        Package("amsmath"),
        Package("amssymb"),
        Package("epsfig"),
        #                   Command('graphicspath{{../}}'),
        #                   Command('frenchspacing'),
        #                   Command('counterwithin{figure}{section}'),
        #                   Command('counterwithin{table}{section}'),
        #                   Command('fancypagestyle{headings}{\\fancyhead{} \\renewcommand{\headrulewidth}{1pt} \\fancyheadoffset{0cm} \\fancyhead[RO]{\\nouppercase{\\leftmark}} \\fancyhead[LE]{\\nouppercase{\\leftmark}} \\fancyfoot{} \\fancyfoot[LE,RO]{\\thepage}}'),
        #                   Command('fancypagestyle{plain}{\\fancyhf{} \\renewcommand{\\headrulewidth}{0pt} \\fancyfoot[LE,RO]{\\thepage}}'),
        #                   Command('numberwithin{equation}{section}'),
        #                   Command('renewcommand{\\familydefault}{\\sfdefault}'),
        # \renewcommand{\familydefault}{\sfdefault}
    ]

    def __init__(
        self,
        default_filepath="default_filepath",
        title="Basic title",
        *,
        documentclass="elsarticle",
        document_options=["final", "times"],  # for submission
        #                  document_options=['final', '5p', 'times' ,'twocolumn'], for preview
        fontenc=None,
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=[
            "inner=30mm",
            "outer=20mm",
            "bindingoffset=10mm",
            "top=25mm",
            "bottom=25mm",
        ],  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        data=None,
    ):

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )

    @property
    def begin_frontmatter(self):
        self.append(ReportText(NoEscape("\\begin{frontmatter}")))

    @property
    def end_frontmatter(self):
        self.append(ReportText(NoEscape("\\end{frontmatter}")))

    def corresponding_author(self, author=None, email=None):

        self.append(
            Command("author[label1]", arguments=[NoEscape(f"{author}\\corref{{cor1}}")])
        )
        self.append(Command("ead", arguments=[f"{email}"]))
        #         self.append(Command('fntext[label2]'))
        self.append(Command("cortext[cor1]{Corresponding author.}"))
        self.append(
            Command(
                "affiliation[label1]",
                arguments=[
                    NoEscape(
                        "organization={Warsaw Univeristy of Technology, Faculty of Automotive and Constrction Machinery Engineering},addressline={Ludwika Narbutta 84},city={Warsaw},postcode={02-524},state={Masovian Voivodeship},country={Poland}"
                    )
                ],
            )
        )

    def keywords(self, keywords_list=None):
        kwl = len(keywords_list)
        with self.create(Keywords()):
            for num, val in enumerate(keywords_list):
                self.append(val)
                if num < kwl - 1:
                    self.append(NoEscape("\sep"))

    def authors(self, authors_list=None):
        self.authors_list = authors_list

        for num, val in enumerate(self.authors_list):
            self.append(Command("author[label1]", arguments=[val]))

    def journal(self, journal_name=None):
        self.preamble.append(Command("journal", arguments=[journal_name]))

    def title(self, title=None):
        self.preamble.append(Command("title", arguments=[title]))

    def abstract(self, container=None):
        with self.create(Abstract()):
            for num, val in enumerate(container):
                self.append(ReportText(val))


class WutThesis(Document):
    """
    A class for creating a Warsaw University of Technology (WUT) thesis document.

    This class extends the `Document` class to include all necessary LaTeX packages, commands, and configurations required for WUT thesis formatting. It provides an easy way to create and manage thesis documents with pre-configured settings.

    Attributes:
        latex_name (str): Name of the document type ('document').
        packages (list): List of LaTeX packages and commands preloaded for thesis formatting.

    Example Usage:
        >>> from dynpy.utilities.documents.document import WutThesis
        >>> from pylatex import Section
        >>> doc = WutThesis('./output/thesis_name')
        >>> section = Section('Introduction')
        >>> doc.append(section)
        >>> doc.generate_pdf(clean_tex=True)

    Example of Full Thesis Workflow:
        >>> from dynpy.utilities.documents.document import WutThesis
        >>> from pylatex import Section, Command
        >>> doc = WutThesis('./output/thesis_name')
        >>> doc.append(Command('title', arguments=['Thesis Title']))
        >>> doc.append(Command('author', arguments=['Student Name']))
        >>> doc.append(Command('date', arguments=['\today']))
        >>> doc.append(Command('maketitle'))
        >>> section = Section('Introduction')
        >>> doc.append(section)
        >>> doc.generate_pdf(clean_tex=True)
    """

    latex_name = "document"
    packages = [
        Package(
            "geometry",
            options=[
                "lmargin=30mm",
                "rmargin=30mm",
                "top=25mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        Package("microtype"),
        Package("authoraftertitle"),
        # Package('polski',options=['MeX']),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Package("graphicx"),
        Package("indentfirst"),
        Package("pdfpages"),
        Package("amsmath"),
        Package("markdown"),
        Package("hyperref"),
        Command("newcommand{\praca}", arguments=["Praca dyplomowa"]),
        Command("newcommand{\dyplom}", arguments=["Magisterska"]),
        Command("newcommand{\kierunek}", arguments=["Wpisać kierunek"]),
        Command("newcommand{\specjalnosc}", arguments=["Wpisać specjalność"]),
        Command("newcommand{\\autor}", arguments=["Imię i nazwisko autora"]),
        Command("newcommand{\opiekun}", arguments=["Wpisać opiekuna"]),
        Command("newcommand{\promotor}", arguments=["Wpisać promotora"]),
        Command("newcommand{\konsultant}", arguments=["Wpisać konsultanta"]),
        Command(
            "newcommand{\\tytul}", arguments=["Wpisać tytuł pracy dyplomowej po polsku"]
        ),
        Command("newcommand{\\album}", arguments=["Wpisać numer albumu"]),
        Command("newcommand{\supervisor}", arguments=["dr inż. Bogumił Chiliński"]),
        Command("newcommand{\\rok}", arguments=["Rok składania pracy"]),
        Command(
            "newcommand{\kluczowe}",
            arguments=["Słowa kluczowe: Wpisać słowa kluczowe po polsku"],
        ),
        # Command('renewcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']),
        Command("graphicspath{{../}}"),
        Command("frenchspacing"),
        Command("counterwithin{figure}{section}"),
        Command("counterwithin{table}{section}"),
        Command(
            "fancypagestyle{headings}{\\fancyhead{} \\renewcommand{\headrulewidth}{1pt} \\fancyheadoffset{0cm} \\fancyhead[RO]{\\nouppercase{\\leftmark}} \\fancyhead[LE]{\\nouppercase{\\leftmark}} \\fancyfoot{} \\fancyfoot[LE,RO]{\\thepage}}"
        ),
        Command(
            "fancypagestyle{plain}{\\fancyhf{} \\renewcommand{\\headrulewidth}{0pt} \\fancyfoot[LE,RO]{\\thepage}}"
        ),
        Command("numberwithin{equation}{section}"),
        Command("renewcommand{\\familydefault}{\\sfdefault}"),
        # \renewcommand{\familydefault}{\sfdefault}
    ]

    def __init__(
        self,
        default_filepath: str = "default_filepath",
        title: str = "Basic title",
        *,
        documentclass: str = "article",
        document_options: list = ["a4paper", "11pt", "twoside"],
        fontenc: str = "T1",
        inputenc: str = "utf8",
        font_size: str = "normalsize",
        lmodern: bool = False,
        textcomp: bool = True,
        microtype: bool = True,
        page_numbers: bool = True,
        indent: Optional[Union[str, int]] = None,
        geometry_options: Optional[list] = [
            "inner=30mm",
            "outer=20mm",
            "bindingoffset=10mm",
            "top=25mm",
            "bottom=25mm",
        ],
        data: Optional[dict] = None,
    ):
        """
        Initialize the WutThesis class with optional customization.

        Args:
            default_filepath (str): Path for the generated document.
            title (str): Title of the document.
            documentclass (str): LaTeX document class (e.g., 'article').
            document_options (list): Options for the document class.
            fontenc (str): Font encoding (default 'T1').
            inputenc (str): Input encoding (default 'utf8').
            font_size (str): Font size (default 'normalsize').
            lmodern (bool): Whether to use Latin Modern fonts.
            textcomp (bool): Whether to use the `textcomp` package.
            microtype (bool): Whether to use the `microtype` package.
            page_numbers (bool): Whether to include page numbers.
            indent (Optional[Union[str, int]]): Indentation settings.
            geometry_options (Optional[list]): Geometry options for page layout.
            data (Optional[dict]): Additional data for customization.
        """

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        #         label=self.label
        self.title = title
        # self.packages.append(Command('title', arguments=[NoEscape(self.title)]))
        #         self.packages.append(Command('date', arguments=[NoEscape('\\today')]))
        #         self.packages.append(Command('newcommand{\praca}', arguments=['Praca dyplomowa']))
        #         self.packages.append(Command('newcommand{\dyplom}', arguments=['Magisterska']))
        #         self.packages.append(Command('newcommand{\kierunek}', arguments=['Wpisać kierunek']))
        #         self.packages.append(Command('newcommand{\specjalnosc}', arguments=['Wpisać specjalność']))
        #         self.packages.append(Command('newcommand{\\autor}', arguments=['Imię i nazwisko autora']))
        #         self.packages.append(Command('newcommand{\opiekun}', arguments=['Wpisać opiekuna']))
        #         self.packages.append(Command('newcommand{\promotor}', arguments=['Wpisać promotora']))
        #         self.packages.append(Command('newcommand{\konsultant}', arguments=['Wpisać konsultanta']))
        #         self.packages.append(Command('newcommand{\\tytul}', arguments=['Wpisać tytuł pracy dyplomowej po polsku']))
        #         self.packages.append(Command('newcommand{\\album}', arguments=['303596']))
        #         self.packages.append(Command('newcommand{\supervisor}', arguments=['dr inż. Bogumił Chiliński']))
        #         self.packages.append(Command('newcommand{\\rok}', arguments=['Rok składania pracy']))
        #         self.packages.append(Command('newcommand{\kluczowe}', arguments=['Słowa kluczowe: Wpisać słowa kluczowe po polsku']))
        #         #self.packages.append(Command('renewcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']))
        self.packages.append(Command("graphicspath{{../}}"))
        #         self.append(Command('maketitle'))
        self.append(NoEscape("%%% New doc"))
        # tu implementować co tam potrzeba

    @classmethod
    def base_setup(cls, create_file: bool = False):
        """
        Sets up the base template for the WUT thesis document.

        This method provides an easy way to initialize a thesis document with the necessary sections, structure, and configuration.

        Args:
            create_file (bool): If True, creates a base setup environment including necessary directories and files.

        Returns:
            Optional[str]: Returns setup content if `create_file` is False, otherwise creates the environment.

        Usage Example:
            >>> from dynpy.utilities.documents.document import WutThesis
            >>> WutThesis.base_setup()

        When to use:
            - Use this method when starting a new WUT thesis document.
            - If `create_file=True`, it will generate all required directories and files for a structured thesis.
            - If `create_file=False`, it returns a string with setup instructions that can be used manually.
        """

        if create_file is True:
            return cls._create_base_setup_env()

        # doc = WutThesis('./output/thesis_name')
        # doc.preamble.append(NoEscape(r'\\usepackage[MeX]{polski}')) #to set polish as main language"""

        thesis_introduction_str = """#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
from dynpy.utilities.report import *

sec_intro = Section('Section that presents text reporting')
CurrentContainer(sec_intro)

sub_problem_outline = Subsection('Outline of the problem')
CurrentContainer(sub_problem_outline)
display(ReportText('This subsection provides information about investigated problem. '*1))

sub_obj_assum = Subsection('Objectives and assumptions')
CurrentContainer(sub_obj_assum)
display(ReportText('This subsection provides objectives and assumptions. '*1))

sub_SOT = Subsection('State of the art')
CurrentContainer(sub_SOT)
display(ReportText('This subsection provides state of the art. '*1))

sub_methodology = Subsection('Methodology')
CurrentContainer(sub_methodology)
display(ReportText('This subsection provides methodology. '*1))"""

        math_str = '''#!!! NO NEEDED RUN PREVIOUS CELLS !!!#
from sympy import Eq, Symbol, symbols
from dynpy.utilities.report import *


sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

sec_description = Subsection('Description of dynamic model')
CurrentContainer(sec_description)
display(ReportText('This subsection provides description of the model. '*1))

from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys

dyn_sys = DynamicSys()
display(dyn_sys.as_picture())

display(ReportText('Summary of the subsection highlighting its major achievements. '*1))

sub_model_code = Subsection('Virtual model of the object')
CurrentContainer(sub_model_code)

from dynpy.utilities.report import ObjectCode

display(ReportText('This subsection provides code of the model. '*1))
display(ObjectCode(DynamicSys))
display(ReportText('Summary of the subsection highlighting its major achievements. '*1))


sec_lagrangian = Subsection('Lagrangian and derivatives')
CurrentContainer(sec_lagrangian)
display(ReportText('This subsection provides calculation of lagrangian and derivatives. '*1))

lagrangian = dyn_sys.L[0]
lagrangian

t=dyn_sys.ivar


for coord in dyn_sys.q:

    display(ReportText(f'Calculations of derivatives for ${vlatex(coord)}$'))


    vel=coord.diff(t)
    diff1 = lagrangian.diff(vel)
    diff2 = diff1.diff(t)
    diff3 = lagrangian.diff(coord)


    diff1_sym = Symbol(f'\\frac{{\\partial L}}{{\\partial {vlatex(vel)}}}')
    diff2_sym = Symbol(f' \\frac{{d }}{{dt}}  {diff1_sym}'  )
    diff3_sym = Symbol(f'\\frac{{\\partial L}}{{\\partial {vlatex(coord)}}}')


    display(SympyFormula(  Eq(diff1_sym,diff1)  ))
    display(SympyFormula(  Eq( diff2_sym ,diff2)  ))
    display(SympyFormula(  Eq( diff3_sym ,  diff3)  ))

display(ReportText('Outcomes of governing equations analysis. '*10))


sec_equations = Subsection('Equations of motion')
CurrentContainer(sec_equations)


display(ReportText('This subsection provides calculation of equations of motion and solution of the system. '*1))

ds1=dyn_sys.eoms
ds1_eqns=ds1.as_eq_list()

for eq1 in ds1.as_eq_list():
    display(SympyFormula(eq1.simplify()))

if ds1.is_solvable():

    ds2=dyn_sys.linearized()._ode_system.general_solution

    for eq2 in ds2.as_eq_list():
        display(SympyFormula(eq2.simplify()))

    ds3=dyn_sys.linearized()._ode_system.steady_solution

    for eq3 in ds3.as_eq_list():
        display(SympyFormula(eq3.simplify()))

    ds4=dyn_sys.linearized().eoms.solution
    ds4_eqns = ds4.as_eq_list()

    for eq4 in ds4_eqns:
        display(SympyFormula(eq4.simplify()))


display(ReportText(f'Outcomes of governing equations {AutoMarker(ds1_eqns[0])} {AutoMarker(ds1_eqns[-1])}) analysis.'))

sec_math_desc = Subsection('Section that contains all symbols descriptions')
CurrentContainer(sec_math_desc)



E_K,F = symbols('E_K,F')
descriptions = {
E_K:r"Kinetic energy",
F:r"Force",
}
syms_dict = descriptions
DescriptionsRegistry().set_descriptions({**syms_dict})
DescriptionsRegistry().reset_registry()
SymbolsDescription.set_default_header('where:')

display(ReportText('Symbols description is as follows: '))


display(SymbolsDescription(expr=E_K*F))


display(ObjectCode("""
# Guide
from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
BasicsOfODESystemGuide();

"""))


# Basics of DynSys usage based simulations are covered in guide to DynSys usage, use the followin call
# from dynpy.utilities.documents.guides import UsageOfDynamicSystemsGuide
# UsageOfDynamicSystemsGuide()
'''

        picture_str = """
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
from dynpy.utilities.report import *

sec_picture = Section('Section that presents pictures reporting')
CurrentContainer(sec_picture)

display(Picture('./dynpy/models/images/taipei101.png',caption = 'Caption of picture'))"""

        simulaion_str = """#!!! NO NEEDED RUNNING PREVIOUS CELLS !!!#


from dynpy.utilities.report import *
from sympy import *
import numpy as np

sec_simulation = Section('Section that contains simulation')
CurrentContainer(sec_simulation)

# Basics of ODESystem based simulations are covered in guide to ODESystem, use the followin call
# from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
# BasicsOfODESystemGuide()

sec_ODESystem = Subsection('ODESystem simulation')
CurrentContainer(sec_ODESystem)

display(ReportText('Firstly create an ODESystem or import it from dynpy.modes.odes.linear.py'))
display(ObjectCode('spring = SpringMassEps.from_reference_data()'))

display(ReportText('Secondly solve the equation using solution method:'))
display(ObjectCode('spring_sol = spring.solution'))

display(ReportText('Last step is data substitution and using numerized method. After that use $compute_solution$ to finalize the simulation:'))

from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys
dyn_sys = DynamicSys()

data_spring = dyn_sys.get_numerical_parameters()
#ds1.subs(data_spring).with_ics([10,0]).numerized().compute_solution(t_span).plot()

from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys
F=DynamicSys.F
g=DynamicSys.g
m=DynamicSys.m
k=DynamicSys.k

slownik ={
F:2,
g:10,
m:5,
k:100,
}

t_span = np.linspace(0.0,1,30)
sym=DynamicSys().eoms.subs(slownik)
wynik= sym.numerized(backend='numpy').compute_solution(t_span,[0.0,0.0])#.plot()
wynik_tab = TimeDataFrame(wynik).to_latex_dataframe().reported(caption='Caption')
display(wynik_tab)
wynik.plot()


# MSM_sim = Subsection('MSM simulation')
# CurrentContainer(MSM_sim)

# display(ReportText('Firstly create an MultiTimeScaleSolution system or import it from dynpy.modes.odes.linear.py'))
# display(ObjectCode('spring = SpringMassMSM.from_reference_data()'))

# display(ReportText('Secondly solve the equation using solution method:'))
# display(ObjectCode('spring_sol = spring.solution'))

# display(ReportText('Last step is data substitution and using numerized method. After that use $compute_solution$ to finalize the simulation:'))
# spring_sol.subs(data_spring).with_ics([10,0]).numerized().compute_solution(t_span).plot()"""

        veryfication_str = """#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

from sympy import *

from sympy.physics import units



from dynpy.utilities.adaptable import *
from dynpy.utilities.report import CurrentContainer, ReportText
import pandas as pd

t = Symbol('t')
power = Symbol('P')
work = Symbol('W')

unit_dict = {
t: units.second,
power: units.watt,
work: units.joule,
}



LatexDataFrame.set_default_units(unit_dict)



sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept. '*200))



### SELECT ONE CASE
## 1st CASE
##import of external data if path "./data_from_ml/test_1.csv" exists

# test_data = './data_from_ml/test_1.csv'
# df = pd.read_csv(test_data, header = None)

# data_tab = TimeDataFrame(df).to_latex_dataframe().reported(caption='Caption')
# graph=data_tab.iloc[:,[1]]

## 2st CASE
## Creating graph from a
# df = pd.DataFrame({t:[0,1,2],'a':[2,3,4],'b':[8,7,6]}).set_index(t)

# df_cols = df.set_axis([power,work],axis=1)

# data_tab = TimeDataFrame(df_cols).to_latex_dataframe().reported(caption='Caption')
# graph = TimeDataFrame(df_cols).to_pylatex_tikz().in_figure(caption='Caption')

# display(data_tab)

# display(ReportText('Obtained data analysis. '*200))

# display(graph)

# display(ReportText('Description of verifications outcomes. '*200))

## 3rd CASE
## dictionary data presentation in DataFrame
# dictionary ={
#     power:300000,
#     work:5000
# }
# display(
#     LatexDataFrame.formatted(
#         data=dictionary, #import data from a dictionary
#         index=['Value'] #set rows title
#     ).map(lambda x: f'${latex(x)}$')
#     .rename_axis('Parameter', axis=1) #set name for a column
#     .transpose()  # This swaps rows and columns, making the DataFrame vertical
#     .reported(caption='A caption for your data frame')
# )

## 4th CASE
## creating a graph based on simulation
# from dynpy.models.mechanics import ForcedSpringMassSystem as DynamicSys
# F=DynamicSys.F
# g=DynamicSys.g
# m=DynamicSys.m
# k=DynamicSys.k
# slownik ={
# F:2,
# g:10,
# m:5,
# k:100,
# }
# t_span = np.linspace(0.0,1,30)
# sym=DynamicSys().eoms.subs(slownik)
# wynik= sym.numerized(backend='numpy').compute_solution(t_span,[0.0,0.0])#.plot()
# wynik_tab = TimeDataFrame(wynik).to_latex_dataframe().reported(caption='Caption')
# display(wynik_tab)
# wynik.plot()

"""

        conclusion_str = """### NO NEEDED RUN OF ALL PREVIOUS CELLS ###

from dynpy.utilities.report import *

sec_conclusion = Section('Section that contains final conclusions')
CurrentContainer(sec_conclusion)

display(ReportText('Conclusions '*200))"""

        symbols_description_str = """#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

from dynpy.utilities.report import *

sec_symbols = Section('Section that contains all symbols descriptions')
CurrentContainer(sec_symbols)

E_K,F = symbols('E_K,F')
descriptions = {
E_K:r"Kinetic energy",
F:r"Force",
}
syms_dict = descriptions
DescriptionsRegistry().set_descriptions({**syms_dict})
DescriptionsRegistry().reset_registry()
SymbolsDescription.set_default_header('  ')

display(SymbolsDescription({**syms_dict}))"""

        document_str = """#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
from dynpy.utilities.documents.document import WutThesis
from dynpy.utilities.report import *

# Creating file
# Be sure *output* folder is in the current directory

thesis_name = './output/report_name' #path for report file

doc = WutThesis(thesis_name)
# Bibliography of quotations
### Select one bibligraphy managment system
####### BibLatex

doc.preamble.append(Package('biblatex',["backend=biber","sorting=none"]))
doc.preamble.append(Command('addbibresource','elementy_bibliagrafia.bib'))
####### Natbib
#doc.preamble.append(Package('natbib')
#doc.preamble.append(Command('bibliographystyle','unsrt'))

## Units
doc.preamble.append(Package('siunitx'))

# TOC

###############################################################################
######### UNCOMMENT IT IF YOU NEED ADD WUT THESIS TITLE PAGES

# doc.append(Command('includepdf{./Images/Front_Page.pdf}')) #includes front page
# doc.append(Command('pagestyle{plain}'))
# doc.append(Command('cleardoublepage'))
# doc.append(Command('includepdf{./Images/oswiadczenie_autora_pracy.pdf}'))
# doc.append(Command('pagestyle{plain}'))
# doc.append(Command('cleardoublepage'))
# doc.append(Command('includepdf{./output/oswiadczenie_biblioteczne.pdf}'))
# doc.append(Command('pagestyle{plain}'))
# doc.append(Command('cleardoublepage'))

###############################################################################
###############################################################################


doc.append(Command('tableofcontents')) #adds TOC

doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)

doc.append(sub_model_code)
doc.append(sec_math_desc)
doc.append(sec_simulation)

doc.append(sec_picture)
doc.append(sec_verification)
#doc.append(sec_tables)
doc.append(sec_symbols)
doc.append(sec_conclusion)


### BibLatex
#doc.append(Command('printbibliography',arguments=["title={Bibliography}"])) - argument is to improve
doc.append(Command('printbibliography',options=[NoEscape("title={Bibliography}")]))

### Natbib
#doc.append(Command('bibliography',arguments=["references"])) # .bib file as "references"


#FIGURES LIST

doc.append(Command('addcontentsline{toc}{section}{List of figures}'))
doc.append(Command('listoffigures'))
doc.append(Command('pagestyle{plain}'))
doc.append(Command('newpage'))

#TABLES LIST

doc.append(Command('addcontentsline{toc}{section}{List of tables}'))
doc.append(Command('renewcommand{\listtablename}{List of tables}'))
doc.append(Command('listoftables'))
doc.append(Command('pagestyle{plain}'))
doc.append(Command('newpage'))

# Generating file
doc.generate_pdf(clean_tex=False)
"""

        #         preliminary_str=(f"""
        # =======
        preliminary_str = f"""

Examplary setup is as follows:

#References to guide imports needed to include - to see the list of available guides insert and run the code below:
```python
from dynpy.utilities.creators import list_of_guides
list_of_guides()
```

## CELL 1
## Imports

```python
{imports_str}
```



## CELL 2
## Thesis introduction

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{thesis_introduction_str}
```



## CELL 3
## Math

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{math_str}
```


## CELL 4
## Picture


```python

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{picture_str}

```


## CELL 5
## Simulation

Basics of ODESystem based simulations are covered in guide to ODESystem, use the following call:

```python
from dynpy.utilities.documents.guides import BasicsOfODESystemGuide,UsageOfDynamicSystemsGuide
BasicsOfODESystemGuide();
```


```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#


{simulaion_str}

```


## CELL 6
## Veryfication

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{veryfication_str}

```

## CELL 7
## Conclusion

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{conclusion_str}

```

## CELL 9
## Symbols description

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{symbols_description_str}

```

## CELL 10
## Document

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

{document_str}
```

"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = f"""
#Example:

#To prepare a simple document with text and images:

#Good practice here is to allocate 1 section per 1 cell

## CELL 1
## Imports

{imports_str}

## CELL 2
## Thesis introduction

{thesis_introduction_str}

## CELL 3
## Math

{math_str}

## CELL 4
## Picture

{picture_str}

## CELL 5
## Simulation

{simulaion_str}

## CELL 6
## Veryfication

{veryfication_str}

## CELL 7
## Conclusion

{conclusion_str}

## CELL 9
## Symbols description

{symbols_description_str}

## CELL 10
## Document

{document_str}

"""

        return ObjectCode(preliminary_str)

    Jupyter_file_content = """{
    "cells": [
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "d3000c",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "from dynpy.utilities.report import *",
                "\\n",
                "from dynpy.utilities.templates.document import WutThesis",
                "\\n",
                "",
                "doc = WutThesis('./output/thesis_name')"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "c3d54b",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "doc.base_setup()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "11531d",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": []
        }
    ],
    "metadata": {
        "kernelspec": {
            "argv": [
                "/usr/bin/python3",
                "-m",
                "ipykernel",
                "--HistoryManager.enabled=False",
                "--matplotlib=inline",
                "-c",
                "%config InlineBackend.figure_formats = set(['retina'])import matplotlib; matplotlib.rcParams['figure.figsize'] = (12, 7)",
                "-f",
                "{connection_file}"
            ],
            "display_name": "Python 3 (system-wide)",
            "env": {},
            "language": "python",
            "metadata": {
                "cocalc": {
                    "description": "Python 3 programming language",
                    "priority": 100,
                    "url": "https://www.python.org/"
                }
            },
            "name": "python3",
            "resource_dir": "/ext/jupyter/kernels/python3"
        },
        "language_info": {
            "codemirror_mode": {
                "name": "ipython",
                "version": 3
            },
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "nbconvert_exporter": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}"""
    articles_bib = """@INPROCEEDINGS{brief_history_of_simulations,
  author={Goldsman, David and Nance, Richard E. and Wilson, James R.},
  booktitle={Proceedings of the 2010 Winter Simulation Conference},
  title={A brief history of simulation revisited},
  year={2010},
  volume={},
  number={},
  pages={567-574},
  doi={10.1109/WSC.2010.5679129}
}
@article{eckhardt1987stan,
  title={Stan Ulam, John von Neumann},
  author={Eckhardt, Roger},
  journal={Los Alamos Science},
  volume={100},
  number={15},
  pages={131},
  year={1987},
  publisher={Los Alamos Scientific Laboratory}
}
@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
@misc{RTPW,
  author = {mgr Tomasz Duda},
  title = {Rysunek Techniczny Podstawowe Wiadomości},
  journal = {},
  year = {},
  number = {},
  pages = {10},
  doi = {}
}
"""

    def _create_directory(path: str) -> None:
        # """
        # Creates a directory at the specified path.

        # Args:
        #     path (str): The directory path to be created. If subdirectories are needed     #     the path should be structured like 'directory/subdirectory'.

        # """

        import os

        try:
            os.makedirs(path)
            print(f"Directory '{path}' created")

        except FileExistsError:
            print(f"Directory '{path}' exists.")

        except PermissionError:
            print(f"Permission denied, you don't have sudo permission")

        except Exception as e:
            print(f"An error occurred: {e}")

    def _create_file(name: str, content: str = None, path: str = None) -> None:
        # """
        #     Creates a file and writes content to it.

        #     Args:
        #         name (str): The name of the file to be created.
        #         content (str, optional): The content to write into the file.
        #         path (str, optional): The directory where the file should be created.
        #                               If specified, it will ensure the directory exists.
        # """
        import os

        if path is not None:
            if not os.path.exists(path):
                _create_directory(path)

            with open(os.path.join(path, name), "w") as file:
                file.write(content)

        else:
            with open(name, "w") as file:
                file.write(content)

        file.close()

    @classmethod
    def _create_base_setup_env(cls) -> None:
        """
        Creates a Jupyter notebook file with WUT thesis base setup and an output directory containing a .bib file.
        """

        Jupyter_file_content = """{
    "cells": [
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "d3000c",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "from dynpy.utilities.report import *",
                "\\n",
                "from dynpy.utilities.templates.document import WutThesis",
                "\\n",
                "",
                "doc = WutThesis('./output/thesis_name')"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "c3d54b",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "doc.base_setup()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "11531d",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": []
        }
    ],
    "metadata": {
        "kernelspec": {
            "argv": [
                "/usr/bin/python3",
                "-m",
                "ipykernel",
                "--HistoryManager.enabled=False",
                "--matplotlib=inline",
                "-c",
                "%config InlineBackend.figure_formats = set(['retina'])import matplotlib; matplotlib.rcParams['figure.figsize'] = (12, 7)",
                "-f",
                "{connection_file}"
            ],
            "display_name": "Python 3 (system-wide)",
            "env": {},
            "language": "python",
            "metadata": {
                "cocalc": {
                    "description": "Python 3 programming language",
                    "priority": 100,
                    "url": "https://www.python.org/"
                }
            },
            "name": "python3",
            "resource_dir": "/ext/jupyter/kernels/python3"
        },
        "language_info": {
            "codemirror_mode": {
                "name": "ipython",
                "version": 3
            },
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "nbconvert_exporter": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}"""
        articles_bib = """@INPROCEEDINGS{brief_history_of_simulations,
  author={Goldsman, David and Nance, Richard E. and Wilson, James R.},
  booktitle={Proceedings of the 2010 Winter Simulation Conference},
  title={A brief history of simulation revisited},
  year={2010},
  volume={},
  number={},
  pages={567-574},
  doi={10.1109/WSC.2010.5679129}
}
@article{eckhardt1987stan,
  title={Stan Ulam, John von Neumann},
  author={Eckhardt, Roger},
  journal={Los Alamos Science},
  volume={100},
  number={15},
  pages={131},
  year={1987},
  publisher={Los Alamos Scientific Laboratory}
}
@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
@misc{RTPW,
  author = {mgr Tomasz Duda},
  title = {Rysunek Techniczny Podstawowe Wiadomości},
  journal = {},
  year = {},
  number = {},
  pages = {10},
  doi = {}
}
"""
        cls._create_file("WUT_Thesis_starter.ipynb", Jupyter_file_content)

        os.makedirs("output")
        os.makedirs("tikzplots")

        cls._create_file("articles.bib", articles_bib, "output")


class ResearchProjectReport(WutThesis):

    @classmethod
    def base_setup(cls):

        preliminary_str = f"""

Examplary setup is as follows:

## CELL 1
## Imports


```python
{imports_str.replace('WutThesis', 'ResearchProjectReport')}
```


## CELL 2
## Project introduction

```python
{thesis_introduction_str_research}
```


## CELL 3
## Math

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

from sympy import Eq, Symbol, symbols

sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

display(ReportText('Mathematical formulas are reported with the support of sympy and it\\'s symbols.'))

a,b = symbols('a b')
display(SympyFormula(Eq(a,b)))
```

## CELL 4
## Picture

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{picture_str}
```

## CELL 5
## Simulation

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{simulaion_str}
```

## CELL 6
## Veryfication

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept '*200))
```

## CELL 7
## Conclusion

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
    {conclusion_str}
```


## CELL 8
## Document

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

# Creating file
# Be sure *output* folder is in the current directory

project_name = './output/report_name' #path for report file

doc = ResearchProjectReport(project_name)
# Bibliography of quotations
doc.preamble.append(NoEscape(r'\\usepackage[backend=bibtex, sorting=none]{{biblatex}}'))
doc.preamble.append(NoEscape(r'\addbibresource{{elementy_bibliagrafia.bib}}'))
doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)
doc.append(sec_picture)
doc.append(sec_simulation)
doc.append(sec_ODESystem)
# doc.append(sec_MSM_sim)
doc.append(sec_verification)
doc.append(sec_conclusion)

# Generating file
doc.generate_pdf(clean_tex=True)
```



"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = f"""

# Examplary setup is as follows:

## CELL 1
## Imports


{imports_str.replace('WutThesis', 'ResearchProjectReport')}


## CELL 2
## Project introduction

{thesis_introduction_str}



## CELL 3
## Math


from sympy import Eq, Symbol, symbols

sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

display(ReportText('Mathematical formulas are reported with the support of sympy and it\\'s symbols.'))

a,b = symbols('a b')
display(SympyFormula(Eq(a,b)))


## CELL 4
## Picture


{picture_str}



## CELL 5
## Simulation

{simulaion_str}

## CELL 6
## Veryfication


sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept '*200))

## CELL 7
## Conclusion

{conclusion_str}

## CELL 8
## Document

# Creating file
# Be sure *output* folder is in the current directory

project_name = './output/report_name' #path for report file

doc = ResearchProjectReport(project_name)
# Bibliography of quotations
doc.preamble.append(NoEscape(r'\\usepackage[backend=bibtex, sorting=none]{{biblatex}}'))
doc.preamble.append(NoEscape(r'\addbibresource{{elementy_bibliagrafia.bib}}'))
doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)
doc.append(sec_picture)
doc.append(sec_simulation)
doc.append(sec_ODESystem)
doc.append(sec_MSM_sim)
doc.append(sec_verification)
doc.append(sec_conclusion)

# Generating file
doc.generate_pdf(clean_tex=True)




"""

        return ObjectCode(preliminary_str)


class DevelopmentProjectReport(WutThesis):

    @classmethod
    def base_setup(cls):

        preliminary_str = f"""

Examplary setup is as follows:

## CELL 1
## Imports


```python
{imports_str.replace('WutThesis', 'DevelopmentProjectReport')}
```


## CELL 2
## Project introduction

```python
{thesis_introduction_str}
```


## CELL 3
## Math

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

from sympy import Eq, Symbol, symbols

sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

display(ReportText('Mathematical formulas are reported with the support of sympy and it\\'s symbols.'))

a,b = symbols('a b')
display(SympyFormula(Eq(a,b)))
```

## CELL 4
## Picture

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{picture_str}
```

## CELL 5
## Simulation

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
{simulaion_str}
```

## CELL 6
## Veryfication

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept '*200))
```

## CELL 7
## Conclusion

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#
    {conclusion_str}
```


## CELL 8
## Document

```python
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

# Creating file
# Be sure *output* folder is in the current directory

project_name = './output/report_name' #path for report file

doc = ResearchProjectReport(project_name)
# Bibliography of quotations
doc.preamble.append(NoEscape(r'\\usepackage[backend=bibtex, sorting=none]{{biblatex}}'))
doc.preamble.append(NoEscape(r'\addbibresource{{elementy_bibliagrafia.bib}}'))
doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)
doc.append(sec_picture)
doc.append(sec_simulation)
doc.append(sec_ODESystem)
# doc.append(sec_MSM_sim)
doc.append(sec_verification)
doc.append(sec_conclusion)

# Generating file
doc.generate_pdf(clean_tex=True)
```



"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = f"""

# Examplary setup is as follows:

## CELL 1
## Imports


{imports_str.replace('WutThesis', 'DevelopmentProjectReport')}


## CELL 2
## Project introduction

{thesis_introduction_str}



## CELL 3
## Math


from sympy import Eq, Symbol, symbols

sec_formula = Section('Section that presents formulas reporting')
CurrentContainer(sec_formula)

display(ReportText('Mathematical formulas are reported with the support of sympy and it\\'s symbols.'))

a,b = symbols('a b')
display(SympyFormula(Eq(a,b)))


## CELL 4
## Picture


{picture_str}



## CELL 5
## Simulation

{simulaion_str}

## CELL 6
## Veryfication


sec_verification = Section('Section that verificates research results')
CurrentContainer(sec_verification)

display(ReportText('Description of verifications concept '*200))

## CELL 7
## Conclusion

{conclusion_str}

## CELL 8
## Document

# Creating file
# Be sure *output* folder is in the current directory

project_name = './output/report_name' #path for report file

doc = ResearchProjectReport(project_name)
# Bibliography of quotations
doc.preamble.append(NoEscape(r'\\usepackage[backend=bibtex, sorting=none]{{biblatex}}'))
doc.preamble.append(NoEscape(r'\addbibresource{{elementy_bibliagrafia.bib}}'))
doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_obj_assum)
doc.append(sub_SOT)
doc.append(sub_methodology)
doc.append(sec_formula)
doc.append(sec_picture)
doc.append(sec_simulation)
doc.append(sec_ODESystem)
doc.append(sec_MSM_sim)
doc.append(sec_verification)
doc.append(sec_conclusion)

# Generating file
doc.generate_pdf(clean_tex=True)




"""

        return ObjectCode(preliminary_str)


class StateOfArtReport(Document):

    latex_name = "document"
    packages = [
        Package(
            "geometry",
            options=[
                "lmargin=30mm",
                "rmargin=30mm",
                "top=25mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        #Package("microtype"),
        Package("authoraftertitle"),
        #Package("polski", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Package("graphicx"),
        Package("indentfirst"),
        Package("pdfpages"),
        Package("amsmath"),
        Command("newcommand{\praca}", arguments=["Praca dyplomowa"]),
        Command("newcommand{\dyplom}", arguments=["Magisterska"]),
        Command("newcommand{\kierunek}", arguments=["Wpisać kierunek"]),
        Command("newcommand{\specjalnosc}", arguments=["Wpisać specjalność"]),
        Command("newcommand{\\autor}", arguments=["Imię i nazwisko autora"]),
        Command("newcommand{\opiekun}", arguments=["Wpisać opiekuna"]),
        Command("newcommand{\promotor}", arguments=["Wpisać promotora"]),
        Command("newcommand{\konsultant}", arguments=["Wpisać konsultanta"]),
        Command(
            "newcommand{\\tytul}", arguments=["Wpisać tytuł pracy dyplomowej po polsku"]
        ),
        Command("newcommand{\\album}", arguments=["Wpisać numer albumu"]),
        Command("newcommand{\supervisor}", arguments=["dr inż. Bogumił Chiliński"]),
        Command("newcommand{\\rok}", arguments=["Rok składania pracy"]),
        Command(
            "newcommand{\kluczowe}",
            arguments=["Słowa kluczowe: Wpisać słowa kluczowe po polsku"],
        ),
        # Command('renewcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']),
        Command("graphicspath{{../}}"),
        Command("frenchspacing"),
        Command("counterwithin{figure}{section}"),
        Command("counterwithin{table}{section}"),
        Command(
            "fancypagestyle{headings}{\\fancyhead{} \\renewcommand{\headrulewidth}{1pt} \\fancyheadoffset{0cm} \\fancyhead[RO]{\\nouppercase{\\leftmark}} \\fancyhead[LE]{\\nouppercase{\\leftmark}} \\fancyfoot{} \\fancyfoot[LE,RO]{\\thepage}}"
        ),
        Command(
            "fancypagestyle{plain}{\\fancyhf{} \\renewcommand{\\headrulewidth}{0pt} \\fancyfoot[LE,RO]{\\thepage}}"
        ),
        Command("numberwithin{equation}{section}"),
        Command("renewcommand{\\familydefault}{\\sfdefault}"),
        # \renewcommand{\familydefault}{\sfdefault}
    ]

    def __init__(
        self,
        default_filepath="default_filepath",
        title="Basic title",
        *,
        documentclass="article",
        document_options=["a4paper", "11pt", "twoside"],
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=[
            "inner=30mm",
            "outer=20mm",
            "bindingoffset=10mm",
            "top=25mm",
            "bottom=25mm",
        ],  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        data=None,
    ):

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        #         label=self.label
        self.title = title
        # self.packages.append(Command('title', arguments=[NoEscape(self.title)]))
        #         self.packages.append(Command('date', arguments=[NoEscape('\\today')]))
        #         self.packages.append(Command('newcommand{\praca}', arguments=['Praca dyplomowa']))
        #         self.packages.append(Command('newcommand{\dyplom}', arguments=['Magisterska']))
        #         self.packages.append(Command('newcommand{\kierunek}', arguments=['Wpisać kierunek']))
        #         self.packages.append(Command('newcommand{\specjalnosc}', arguments=['Wpisać specjalność']))
        #         self.packages.append(Command('newcommand{\\autor}', arguments=['Imię i nazwisko autora']))
        #         self.packages.append(Command('newcommand{\opiekun}', arguments=['Wpisać opiekuna']))
        #         self.packages.append(Command('newcommand{\promotor}', arguments=['Wpisać promotora']))
        #         self.packages.append(Command('newcommand{\konsultant}', arguments=['Wpisać konsultanta']))
        #         self.packages.append(Command('newcommand{\\tytul}', arguments=['Wpisać tytuł pracy dyplomowej po polsku']))
        #         self.packages.append(Command('newcommand{\\album}', arguments=['303596']))
        #         self.packages.append(Command('newcommand{\supervisor}', arguments=['dr inż. Bogumił Chiliński']))
        #         self.packages.append(Command('newcommand{\\rok}', arguments=['Rok składania pracy']))
        #         self.packages.append(Command('newcommand{\kluczowe}', arguments=['Słowa kluczowe: Wpisać słowa kluczowe po polsku']))
        #         #self.packages.append(Command('renewcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']))
        self.packages.append(Command("graphicspath{{../}}"))
        #         self.append(Command('maketitle'))
        self.append(NoEscape("%%% New doc"))
        # tu implementować co tam potrzeba

    @classmethod
    def base_setup(cls, create_file=False):

        if create_file is True:
            return cls._create_base_setup_env()

        preliminary_str = """

Examplary setup is as follows:

#References to guide imports needed to include - to see the list of available guides insert and run the code below:
```{python}
from dynpy.utilities.creators import list_of_guides
list_of_guides()
```

## CELL 1
## Imports

```{python}
#Create file output
#Create file Images
#In file output create bibliography as .bib file (biblio.bib)

'''# File content begin
@misc{DynPi,
author={GitHub},
title="bogumilchilinski/dynpy",
url={https://github.com/bogumilchilinski/dynpy},
note="Accessed:2024-05-04"
}

@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
'''# File content end



from dynpy.utilities.report import *
from dynpy.utilities.documents.document import StateOfArtReport


from dynpy.utilities.adaptable import TimeDataFrame
import pandas as pd


biblio_mngr=BiblographyManager(arguments='biblio.bib')

biblio_entries=NoEscape(
r'''
@misc{DynPi,
author={GitHub},
title="bogumilchilinski/dynpy",
url={https://github.com/bogumilchilinski/dynpy},
note="Accessed:2024-05-04"
}

@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
''')

biblio_mngr.append(biblio_entries)


doc = StateOfArtReport('./output/report_name')
```



## CELL 2
## Document introduction

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_intro = Section('Section that presents text reporting')
CurrentContainer(sec_intro)

sub_problem_outline = Subsection('Outline of the problem')
CurrentContainer(sub_problem_outline)
display(ReportText('This subsection provides information about investigated problem. '*100))

sub_SOT = Subsection('State of the art')
CurrentContainer(sub_SOT)
display(ReportText('This subsection provides state of the art with multiple reference for \\cite{DynPi}. '*50))
display(Markdown('Remeber about `biblio.bib` file with referencens. Place it in working directory and subfolder with output files. '*50))

sub_conclusions = Subsection('Conclusions')
CurrentContainer(sub_conclusions)
display(ReportText('This subsection provides methodology. '*100))
```
## CELL 3
## Conclusion

    #!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
    #!!!       BECAUSE OF NEEDED IMPORTS    !!!#

    sec_conclusion = Section('Section that contains final conclusions')
    CurrentContainer(sec_conclusion)

    display(ReportText('Conclusions '*200))

## CELL 4
## Document

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

# Creating file
# Be sure *output* folder is in the current directory

sota_name = './output/report_name' #path for report file

doc = StateOfArtReport(sota_name)
doc.preamble.append(biblio_mngr)

# Bibliography of quotations
### Select one bibligraphy managment system
####### BibLatex
#doc.preamble.append(Package('biblatex',["backend=biber","sorting=none"]))
#doc.preamble.append(Command('addbibresource','biblio.bib'))

####### Natbib
doc.preamble.append(Package('natbib')
doc.preamble.append(Command('bibliographystyle','unsrt'))

# TOC
doc.append(Command('tableofcontents')) #adds TOC

doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_SOT)
doc.append(sub_conclusions)
doc.append(sec_conclusion)


### BibLatex
#doc.append(Command('printbibliography',options=[NoEscape("title={Bibliography}")]))

### Natbib
doc.append(Command('bibliography',arguments=["references"])) # .bib file as "references"

# Generating file
doc.generate_pdf(clean_tex=False)
```

"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = """

#Example:

#To prepare a simple document with text and images:

#Good practice here is to allocate 1 section per 1 cell



## CELL 1
## Imports


#Create file output
#Create file Images
#In file output create bibliography as .bib file (biblio.bib)

'''# File content begin
@misc{DynPi,
author={GitHub},
title="bogumilchilinski/dynpy",
url={https://github.com/bogumilchilinski/dynpy},
note="Accessed:2024-05-04"
}

@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
'''# File content end



from dynpy.utilities.report import *
from dynpy.utilities.documents.document import StateOfArtReport


from dynpy.utilities.adaptable import TimeDataFrame
import pandas as pd


biblio_mngr=BibliographyManager(arguments='biblio.bib')

biblio_entries=NoEscape(
r'''
@misc{
  DynPi,
  author={GitHub},
  title="bogumilchilinski/dynpy",
  url={https://github.com/bogumilchilinski/dynpy},
  note="Accessed:2024-05-04"
}

@book{
  lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}

@misc{
  NumPy,
  author="NumPy",
  title={https://numpy.org/},
  url={https://numpy.org/}, journal={NumPy}
}
''')

biblio_mngr.append(biblio_entries)


doc = StateOfArtReport('./output/report_name')


## CELL 2
## Document introduction


#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_intro = Section('Section that presents text reporting')
CurrentContainer(sec_intro)

sub_problem_outline = Subsection('Outline of the problem')
CurrentContainer(sub_problem_outline)
display(ReportText('This subsection provides information about investigated problem.'))

sub_SOT = Subsection('State of the art')
CurrentContainer(sub_SOT)
display(ReportText('This subsection provides state of the art with multiple reference for \\cite{DynPi}.'))


sub_conclusions = Subsection('Conclusions')
CurrentContainer(sub_conclusions)
display(ReportText('This subsection provides methodology \\cite{NumPy}.'))


## CELL 3
## Conclusion

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_conclusion = Section('Section that contains final conclusions')
CurrentContainer(sec_conclusion)

display(ReportText('Conclusions'))

## CELL 4
## Document


#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

# Creating file
# Be sure *output* folder is in the current directory

sota_name = './output/report_name' #path for report file

doc = StateOfArtReport(sota_name)
doc.preamble.append(biblio_mngr)

# Bibliography of quotations
### Select one bibligraphy managment system
####### BibLatex
#doc.preamble.append(Package('biblatex',["backend=biber","sorting=none"]))
#doc.preamble.append(Command('addbibresource','biblio.bib'))

####### Natbib
doc.preamble.append(Package('natbib'))
doc.preamble.append(Command('bibliographystyle','unsrt'))

# TOC
doc.append(Command('tableofcontents')) #adds TOC

doc.append(sec_intro) # adding certain sections
doc.append(sub_problem_outline)
doc.append(sub_SOT)
doc.append(sub_conclusions)
doc.append(sec_conclusion)


### BibLatex
#doc.append(Command('printbibliography',options=[NoEscape("title={Bibliography}")]))

### Natbib
doc.append(Command('bibliography',arguments=["biblio"])) # .bib file as "biblio"

# Generating file
doc.generate_pdf(clean_tex=False)


"""
        return ObjectCode(preliminary_str)

    Jupyter_file_content = """{
    "cells": [
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "d3000c",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "from dynpy.utilities.report import *",
                "\\n",
                "from dynpy.utilities.templates.document import WutThesis",
                "\\n",
                "",
                "doc = WutThesis('./output/thesis_name')"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "c3d54b",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "doc.base_setup()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "11531d",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": []
        }
    ],
    "metadata": {
        "kernelspec": {
            "argv": [
                "/usr/bin/python3",
                "-m",
                "ipykernel",
                "--HistoryManager.enabled=False",
                "--matplotlib=inline",
                "-c",
                "%config InlineBackend.figure_formats = set(['retina'])import matplotlib; matplotlib.rcParams['figure.figsize'] = (12, 7)",
                "-f",
                "{connection_file}"
            ],
            "display_name": "Python 3 (system-wide)",
            "env": {},
            "language": "python",
            "metadata": {
                "cocalc": {
                    "description": "Python 3 programming language",
                    "priority": 100,
                    "url": "https://www.python.org/"
                }
            },
            "name": "python3",
            "resource_dir": "/ext/jupyter/kernels/python3"
        },
        "language_info": {
            "codemirror_mode": {
                "name": "ipython",
                "version": 3
            },
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "nbconvert_exporter": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}"""
    articles_bib = """@INPROCEEDINGS{brief_history_of_simulations,
  author={Goldsman, David and Nance, Richard E. and Wilson, James R.},
  booktitle={Proceedings of the 2010 Winter Simulation Conference},
  title={A brief history of simulation revisited},
  year={2010},
  volume={},
  number={},
  pages={567-574},
  doi={10.1109/WSC.2010.5679129}
}
@article{eckhardt1987stan,
  title={Stan Ulam, John von Neumann},
  author={Eckhardt, Roger},
  journal={Los Alamos Science},
  volume={100},
  number={15},
  pages={131},
  year={1987},
  publisher={Los Alamos Scientific Laboratory}
}
@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
@misc{RTPW,
  author = {mgr Tomasz Duda},
  title = {Rysunek Techniczny Podstawowe Wiadomości},
  journal = {},
  year = {},
  number = {},
  pages = {10},
  doi = {}
}
"""

    def _create_directory(path):
        """
        Method create directory from given path

        Arguments:
        path - (str), name of directory that method will create, if subdirectories shall be created, path should look like this: 'directory/subdirectory'
        """
        import os

        try:
            os.makedirs(path)
            print(f"Directory '{path}' created")

        except FileExistsError:
            print(f"Directory '{path}' exists.")

        except PermissionError:
            print(f"Permission denied, you don't have sudo permission")

        except Exception as e:
            print(f"An error occurred: {e}")

    def _create_file(name, content=None, path=None):
        """
        Method create file and write content to it

        Arguments:
        name - (str), name of file created by method,
        content - (str), optional,  string with a content of file, that is going to be writen to it,
        path - (str), optional, directory where file should be created
        """
        import os

        if path is not None:
            if not os.path.exists(path):
                _create_directory(path)

            with open(os.path.join(path, name), "w") as file:
                file.write(content)

        else:
            with open(name, "w") as file:
                file.write(content)

        file.close()

    @classmethod
    def _create_base_setup_env(cls):
        """
        Method that create Jupyter notebook file with WUT thesis base setup and output directory with .bib file
        """

        Jupyter_file_content = """{
    "cells": [
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "d3000c",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "from dynpy.utilities.report import *",
                "\\n",
                "from dynpy.utilities.templates.document import WutThesis",
                "\\n",
                "",
                "doc = WutThesis('./output/thesis_name')"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "c3d54b",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "doc.base_setup()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "11531d",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": []
        }
    ],
    "metadata": {
        "kernelspec": {
            "argv": [
                "/usr/bin/python3",
                "-m",
                "ipykernel",
                "--HistoryManager.enabled=False",
                "--matplotlib=inline",
                "-c",
                "%config InlineBackend.figure_formats = set(['retina'])import matplotlib; matplotlib.rcParams['figure.figsize'] = (12, 7)",
                "-f",
                "{connection_file}"
            ],
            "display_name": "Python 3 (system-wide)",
            "env": {},
            "language": "python",
            "metadata": {
                "cocalc": {
                    "description": "Python 3 programming language",
                    "priority": 100,
                    "url": "https://www.python.org/"
                }
            },
            "name": "python3",
            "resource_dir": "/ext/jupyter/kernels/python3"
        },
        "language_info": {
            "codemirror_mode": {
                "name": "ipython",
                "version": 3
            },
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "nbconvert_exporter": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}"""
        articles_bib = """@INPROCEEDINGS{brief_history_of_simulations,
  author={Goldsman, David and Nance, Richard E. and Wilson, James R.},
  booktitle={Proceedings of the 2010 Winter Simulation Conference},
  title={A brief history of simulation revisited},
  year={2010},
  volume={},
  number={},
  pages={567-574},
  doi={10.1109/WSC.2010.5679129}
}
@article{eckhardt1987stan,
  title={Stan Ulam, John von Neumann},
  author={Eckhardt, Roger},
  journal={Los Alamos Science},
  volume={100},
  number={15},
  pages={131},
  year={1987},
  publisher={Los Alamos Scientific Laboratory}
}
@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
@misc{RTPW,
  author = {mgr Tomasz Duda},
  title = {Rysunek Techniczny Podstawowe Wiadomości},
  journal = {},
  year = {},
  number = {},
  pages = {10},
  doi = {}
}
"""
        cls._create_file("WUT_Thesis_starter.ipynb", Jupyter_file_content)

        os.makedirs("output")
        os.makedirs("tikzplots")

        cls._create_file("articles.bib", articles_bib, "output")

class Report(StateOfArtReport):
    pass
class ThesisTemplate(WutThesis):
    pass


class MechanicalCase(Guide):

    latex_name = "document"
    packages = [
        Package(
            "geometry",
            options=[
                "lmargin=25mm",
                "rmargin=25mm",
                "top=30mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        Package("microtype"),
        Package("authoraftertitle"),
        Package("polski", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Command("pagestyle", arguments=["fancy"]),
        Command("fancyhf", arguments=[""]),
        Command("fancyhead", arguments=["DynPy Team"], options=["R"]),
        Command("fancyhead", arguments=["Mechanics, 2023"], options=["L"]),
        Command("fancyfoot", arguments=[NoEscape("\\thepage")], options=["C"]),
    ]


class CaseStudy(Guide):

    latex_name = "document"
    _documentclass = "article"
    packages = [
        Package(
            "geometry",
            options=[
                "lmargin=25mm",
                "rmargin=25mm",
                "top=30mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        Package("microtype"),
        Package("authoraftertitle"),
        Package("polski", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Command("pagestyle", arguments=["fancy"]),
        Command("fancyhf", arguments=[""]),
        Command("fancyhead", arguments=["B. Chiliński"], options=["R"]),
        Command("fancyhead", arguments=["Studium przypadku, 2023"], options=["L"]),
        Command("fancyfoot", arguments=[NoEscape("\\thepage")], options=["C"]),
    ]


class TechThriveMechanicalCase(Guide):

    latex_name = "document"
    _documentclass = "article"
    packages = [
        Package("float"),
        Package("graphicx"),
        Command("graphicspath{{../}}"),
        Package(
            "geometry",
            options=[
                "lmargin=25mm",
                "rmargin=25mm",
                "top=30mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        Package("microtype"),
        Package("authoraftertitle"),
        Package("polski", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Package("svg"),
        Command("pagestyle", arguments=["fancy"]),
        Command("fancyhf", arguments=[""]),
        Command(
            "fancyhead",
            arguments=[
                NoEscape(
                    r"\includegraphics[height=0.8cm]{./dynpy/models/images/TT.png}"
                )
            ],
            options=["C"],
        ),  #
        Command("fancyhead", arguments=["Mechanika Ogólna"], options=["L"]),
        Command("fancyhead", arguments=[NoEscape("\\today")], options=["R"]),
        Command("fancyfoot", arguments=[NoEscape("\\thepage")], options=["C"]),
        Command("fancyfoot", arguments=["Copyright TechThrive 2023"], options=["L"]),
        Command("fancyfoot", arguments=["Powered by DynPy"], options=["R"]),
        Command("graphicspath{{../}}"),
    ]


class TechThriveODECase(TechThriveMechanicalCase):
    packages = [
        Command("fancyhead", arguments=["Równania różniczkowe"], options=["L"]),
    ]


class TechThriveMVCase(TechThriveMechanicalCase):

    packages = [
        Command("fancyhead", arguments=["Drgania mechaniczne"], options=["L"]),
    ]

    @property
    def default_reported_object(self):

        from ...models.mechanics import (
            ForcedDampedTrolleysWithSprings,
            ForcedSpringMassSystem,
        )

        return ForcedSpringMassSystem()

    @property
    def _report_components(self):

        comp_list = [
            #         mech_comp.TitlePageComponent,
            #         #mech_comp.SchemeComponent,
            #         #mech_comp.ExemplaryPictureComponent,
            #         mech_comp.KineticEnergyComponent,
            #         mech_comp.PotentialEnergyComponent,
            #         mech_comp.LagrangianComponent,
            #         mech_comp.GoverningEquationComponent,
            #         #mech_comp.FundamentalMatrixComponent,
            #         mech_comp.GeneralSolutionComponent,
            #         #mech_comp.SteadySolutionComponent,
        ]

        return comp_list


class QuotationTemplate(Document):
    packages = [
        Package("geometry", options=["left=1in", "top=1in", "right=1in", "bottom=1in"]),
        Package("inputenc", options=["utf8"]),
        Package("hyperref", options=["colorlinks"]),
        Package("graphicx"),
        Package("tabularx"),
        Package("multirow"),
        Package("ragged2e"),
        Package("hhline"),
        Package("array"),
        Package("booktabs"),
        Package("polski", options=["MeX"]),
    ]

    def __init__(
        self,
        default_filepath="default_filepath",
        title="Wycena zlecenia",
        *,
        documentclass="report",
        document_options=None,
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,  # ['lmargin=25mm', 'rmargin=25mm',  'top=20mm', 'bmargin=25mm', 'headheight=50mm'],
        data=None,
    ):

        super().__init__(
            #             title=title,
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        #         label=self.label
        #         self.title=title
        self.preamble.append(Command("hypersetup", "urlcolor=blue"))
        self.preamble.append(
            Command(
                "newcolumntype",
                arguments="R",
                options="1",
                extra_arguments=NoEscape(
                    ">{\\raggedleft\\let\\newline\\\\\\arraybackslash\\hspace{0pt}}m{#1}"
                ),
            )
        )


#         self.packages.append(Command('date', arguments=[NoEscape('\\today')]))
# #         self.append(Command('maketitle'))
#         self.append(NewPage())
# tu implementować co tam potrzeba


class QuotationTemplate(Document):
    packages = [
        Package("geometry", options=["left=1cm", "top=1cm", "right=1cm", "bottom=1cm"]),
        Package("inputenc", options=["utf8"]),
        Package("hyperref", options=["colorlinks"]),
        Package("graphicx"),
        Package("tabularx"),
        Package("multirow"),
        Package("ragged2e"),
        Package("hhline"),
        Package("array"),
        Package("booktabs"),
        Package("polski", options=["MeX"]),
    ]

    def __init__(
        self,
        default_filepath="default_filepath",
        title="Wycena zlecenia",
        *,
        documentclass="report",
        document_options=None,
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,  # ['lmargin=25mm', 'rmargin=25mm',  'top=20mm', 'bmargin=25mm', 'headheight=50mm'],
        data=None,
    ):

        super().__init__(
            #             title=title,
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        #         label=self.label
        #         self.title=title
        self.preamble.append(Command("hypersetup", "urlcolor=blue"))
        self.preamble.append(
            Command(
                "newcolumntype",
                arguments="R",
                options="1",
                extra_arguments=NoEscape(
                    ">{\\raggedleft\\let\\newline\\\\\\arraybackslash\\hspace{0pt}}m{#1}"
                ),
            )
        )


#         self.packages.append(Command('date', arguments=[NoEscape('\\today')]))
# #         self.append(Command('maketitle'))
#         self.append(NewPage())
# tu implementować co tam potrzeba


class ScheduleTemplate(QuotationTemplate):
    pass


class LandscapeScheduleTemplate(ScheduleTemplate):
    def __init__(
        self,
        default_filepath="default_filepath",
        title="Wycena zlecenia",
        *,
        documentclass="report",
        document_options="landscape",
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,  # ['lmargin=25mm', 'rmargin=25mm',  'top=20mm', 'bmargin=25mm', 'headheight=50mm'],
        data=None,
    ):

        super().__init__(
            #             title=title,
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        #         label=self.label
        #         self.title=title
        self.preamble.append(Command("hypersetup", "urlcolor=blue"))
        self.preamble.append(
            Command(
                "newcolumntype",
                arguments="R",
                options="1",
                extra_arguments=NoEscape(
                    ">{\\raggedleft\\let\\newline\\\\\\arraybackslash\\hspace{0pt}}m{#1}"
                ),
            )
        )


class BeamerPresentation(Document):

    latex_name = "document"
    packages = [
        #Package("microtype"),
        # Package('polski',options=['MeX']),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        # Package('titlesec'),
        # Package('fancyhdr'),
        # Command('pagestyle', arguments=['fancy']),
        Command("author", arguments=["Author name"]),
        # Command('fancyhead', arguments=[NoEscape('\includegraphics[height=1.5cm]{./images/logoPOWER.jpg}')],options=['C']),
        # Command('fancyfoot', arguments=['BCh&KT'],options=['R']),
        # Command('fancyfoot', arguments=['Practical Python, 2022'],options=['L']),
        # Command('fancyfoot', arguments=[NoEscape('\\thepage')],options=['C']),
        Command("usetheme", arguments=["Madrid"]),
        Command("graphicspath", arguments=[NoEscape("{../}")]),
    ]

    @classmethod
    def base_setup(cls):

        preliminary_str = """

# Examplary setup is as follows:

from dynpy.utilities.documents.document import BeamerPresentation
from dynpy.utilities.report import*
SympyFormula._break_mode = 'eq' # neccesarry for correct equation display


#CELL_1
frame_1=Frame(title='Introduction',options=[])
CurrentContainer(frame_1)
display(ReportText('abc '*100))


#CELL_2
frame_2=Frame(title='Physical model',options=['allowframebreaks'])
CurrentContainer(frame_2)
display(ReportText('abc '*100))


#CELL_3
from sympy import *

frame_3=CodeFrame(title='Simulation results',options=[])
CurrentContainer(frame_3)

display(ObjectCode('from dynpy.utilities.report import *'))

#CELL_4
from sympy import *

frame_4=Frame(title='Simulation results',options=[])
CurrentContainer(frame_4)

display(SympyFormula(Symbol('a')**2))

#CELL_5
doc = BeamerPresentation('Example_beamer',title='Exemplary presentation')
doc.append(frame_1)
doc.append(frame_2)
doc.append(frame_3)
doc.append(frame_4)
doc.generate_pdf(clean_tex=False)


"""
        return ObjectCode(preliminary_str)

    def __init__(
        self,
        default_filepath="default_filepath",
        *,
        title=None,
        author=None,
        initials=None,
        documentclass="beamer",
        document_options=None,
        fontenc="T1",
        inputenc="utf8",
        font_size="footnotesize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,
        data=None,
    ):

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        self.title = title
        self.author = author
        self.initials = initials

        if self.title is not None:
            self.packages.append(Command("title", arguments=[self.title], options=[""]))

        if self.author is not None:
            self.packages.remove(Command("author", arguments=["Author name"]))

            if self.initials is not None:
                
                self.packages.append(
                    Command("author", arguments=[self.author], options=[self.initials])
                )
            else:
                self.packages.append(
                    Command("author", arguments=[self.author], options=[""])
                )

        self.append(Command("frame", arguments=[NoEscape(r"\titlepage")]))


class BeamerTemplate(BeamerPresentation):

    pass


class DGBeamer(Document):

    latex_name = "document"
    packages = [
        Package("microtype"),
        # Package('english',options=['MeX']),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        # Package('titlesec'),
        # Package('fancyhdr'),
        # Command('pagestyle', arguments=['fancy']),
        Command("author", arguments=["Author name"]),
        # Command('fancyhead', arguments=[NoEscape('\includegraphics[height=1.5cm]{./images/logoPOWER.jpg}')],options=['C']),
        # Command('fancyfoot', arguments=['BCh&KT'],options=['R']),
        # Command('fancyfoot', arguments=['Practical Python, 2022'],options=['L']),
        # Command('fancyfoot', arguments=[NoEscape('\\thepage')],options=['C']),
        Command("usetheme", arguments=["Madrid"]),
        Command("graphicspath", arguments=[NoEscape("{../}")]),
    ]

    def __init__(
        self,
        default_filepath="default_filepath",
        *,
        title=None,
        documentclass="beamer",
        document_options=None,
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,
        data=None,
    ):

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        self.title = title

        if self.title is not None:
            self.packages.append(Command("title", arguments=[self.title]))
        self.append(Command("frame", arguments=[NoEscape(r"\titlepage")]))


class PosterTemplate(Document):

    latex_name = "document"
    packages = [
        Package("microtype"),
        Package("polski", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("siunitx"),
        # Package('titlesec'),
        # Package('fancyhdr'),
        # Command('pagestyle', arguments=['fancy']),
        Command("author", arguments=["Author name"]),
        # Command('fancyhead', arguments=[NoEscape('\includegraphics[height=1.5cm]{./images/logoPOWER.jpg}')],options=['C']),
        # Command('fancyfoot', arguments=['BCh&KT'],options=['R']),
        # Command('fancyfoot', arguments=['Practical Python, 2022'],options=['L']),
        # Command('fancyfoot', arguments=[NoEscape('\\thepage')],options=['C']),
        Command("usetheme", arguments=["Simple"]),
        Command(
            "institute",
            arguments=[
                "Institute of Machine Design Fundamentals, Warsaw University Of Technology"
            ],
        ),
        Command("graphicspath", arguments=[NoEscape("{../}")]),
    ]

    @classmethod
    def base_setup(cls):

        preliminary_str = """

# Examplary setup is as follows:

from dynpy.utilities.documents.document import PosterTemplate
from dynpy.utilities.report import*
from pylatex.base_classes import Environment, ContainerCommand
SympyFormula._break_mode = 'eq'
class PosterBlock(ContainerCommand):
    latex_name='block'

#CELL_1
Introduction = PosterBlock('Introduction')
CurrentContainer(Introduction)
display(ReportText('abc '*100))

#CELL_2
Motivation=PosterBlock('Motivation for the Research')
CurrentContainer(Motivation)
display(ReportText('abc '*100))

#CELL_3
from sympy import *

Simulation=PosterBlock('Simulation results')
CurrentContainer(Simulation)

display(SympyFormula(Symbol('a')**2))

#CELL_4
Physical=PosterBlock('Physical model')
CurrentContainer(Physical)
display(ReportText('abc '*100))



#CELL_5
doc = PosterTemplate('./output/template',title=r'{Title of the paper}',document_options=['a1paper','landscape', '12pt'])
doc.preamble.append(NoEscape(r'\author{authors}'))
doc.append(NoEscape(r'\begin{columns}'))
doc.append(NoEscape(r'\column{0.33}'))
doc.append(Introduction)
doc.append(NoEscape(r'\column{0.33}'))
doc.append(Motivation)
doc.append(Simulation)
doc.append(NoEscape(r'\column{0.33}'))
doc.append(Physical)

doc.append(NoEscape(r'\end{columns}'))
doc.generate_pdf('./output/template',clean_tex=False)






"""
        return ObjectCode(preliminary_str)

    def __init__(
        self,
        default_filepath="default_filepath",
        *,
        title=None,
        documentclass="tikzposter",
        document_options=None,
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=None,
        data=None,
    ):

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        self.title = title
        title_with_box = NoEscape(
            r"\parbox{0.9\linewidth}{\centering " + self.title + " }"
        )

        if self.title is not None:
            self.packages.append(Command("title", arguments=[title_with_box]))

        # self.append(Command('frame', arguments=[NoEscape(r'\titlepage')]))
        self.append(Command("maketitle"))


class MDPIPaper(Document):

    latex_name = "document"
    packages = []
    title = None
    cwd = os.getcwd()
    path1 = f"{cwd}/Definitions"

    def __init__(
        self,
        default_filepath="default_filepath",
        title=None,
        documentclass=None,
        journal=None,
        document_options=None,  # for submission
        fontenc=None,
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=False,
        microtype=False,
        page_numbers=True,
        indent=None,
        geometry_options=None,  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        data=None,
    ):

        if document_options is None:
            document_options = [journal, "article", "submit", "pdftex", "moreauthors"]

        super().__init__(
            default_filepath=default_filepath,
            documentclass=["Definitions/mdpi"],
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        

        Markdown.set_mdpi() #deletes hyperref package associated with Markdown class to avoid clash
        
        self.preamble.append(Command('firstpage','1'))
        self.preamble.append(Command('makeatletter'))
        self.preamble.append(NoEscape('\\setcounter{page}{\@firstpage}'))
        self.preamble.append(Command('makeatother'))
        self.preamble.append(Command('pubvolume','1'))
        self.preamble.append(Command('issuenum','1'))
        self.preamble.append(Command('articlenumber','0'))
        self.preamble.append(Command('pubyear','2025'))
        self.preamble.append(Command('copyrightyear','2025'))
        self.preamble.append(Command('datereceived',' '))
        self.preamble.append(Command('daterevised',' '))
        self.preamble.append(Command('dateaccepted',' '))
        self.preamble.append(Command('datepublished',' '))
        self.preamble.append(Command('hreflink','https://doi.org/'))
        self.preamble.append(Command('Title',title))
        self.preamble.append(Command('TitleCitation',title))
        self.preamble.append(Command('firstnote','Current address: Affiliation'))
        self.preamble.append(Command('secondnote','These authors contributed equally to this work.'))
        #self.append(NewPage())
        # tu implementować co tam potrzeba

        cwd = os.getcwd()

        source_path = (
            f"/home/user/Shared files/modules/dynpy/utilities/documents/Definitions"
        )
        path1 = f"{cwd}/Definitions"
        path2 = f"{cwd}/output/Definitions"

        if os.path.exists(path1) == False:
            shutil.copytree(source_path, path1)
        if os.path.exists(path2) == False:
            shutil.copytree(source_path, path2)

    def authors(self, orcid_dict, corr_no=1):

        ######## should be edited further if more affiliations!
        # for orcid_dict provide: consecutive capital latin alphabet letters : [orcid_number,firstnameandlastname]; {A:[00001,'Damian Sierociński'],B:[000002,'Bogumił Chiliński'],...}
        # corr_no: which of the authors in provided list is the corresponding author; int; count starts from 1
        affiliation_no = 1
        author_no = 1
        authors_no = len(orcid_dict)
        author_string = "\\Author{"
        authornames_string = "\\AuthorNames{"
        authorcitation_string = "\\AuthorCitation{"
        self.od = orcid_dict
        for key, val in orcid_dict.items():
            display(key)
            self.preamble.append(
                NoEscape("\\newcommand{\\orcidauthor" + key + "}{" + val[0] + "}")
            )

            if len(val) > 2:
                affiliation_no = val[2]

            author_string += (
                val[1] + " $^{" + str(affiliation_no) + ",\\dagger,\\ddagger"
            )
            authornames_string += val[1]
            if author_no == corr_no:
                author_string += ",*"
            author_string += "}$\\orcid" + key + "{}"
            if author_no != authors_no and author_no != authors_no - 1:
                author_string += ", "
                authornames_string += ", "
            elif author_no == authors_no - 1:
                author_string += " and "
                authornames_string += " and "

            author_no += 1

        self.preamble.append(NoEscape(author_string + "}"))
        self.preamble.append(NoEscape(authornames_string + "}"))
        self.preamble.append(
            NoEscape(authorcitation_string + self._citation_format() + "}")
        )

    def _citation_format(self):
        display(self.od)
        cit_string = ""
        for key, val in self.od.items():
            firstname = val[1].split(" ")[0]
            firstname = firstname[0] + "."
            lastname = val[1].split(" ")[1]
            display(firstname)
            display(lastname)
            cit_string += lastname + " " + firstname + "; "

        return cit_string[:-2]

    def address(self):
        ######## should be edited further if more affiliations!
        self.preamble.append(
            NoEscape(
                "\\address{$^{1}$ \\quad Department of Computer Techniques, Institute of Machine Design Fundamentals, Faculty of Automotive and Construction Machinery Engineering, Warsaw University of Technology; 84 Ludwika Narbutta Street, 02-524 Warsaw, Poland}"
            )
        )

    def correspondence(self, email):
        self.preamble.append(NoEscape("\\corres{Correspondence: " + email + "}"))

    def abstract(self, text):
        self.preamble.append(Command("abstract", text))

    def keywords(self, keywords_list):

        kwl_string = "\\keyword{"
        key_num = 0
        for num, val in enumerate(keywords_list):
            kwl_string += val
            if key_num != len(keywords_list) - 1:
                kwl_string += "; "
            key_num += 1
        self.preamble.append(NoEscape(kwl_string + "}"))


#             \AuthorCitation{Lastname, F.; Lastname, F.; Lastname, F.}
class RevievPrototype(Document):

    latex_name = "document"
    packages = [
        Package(
            "geometry",
            options=[
                "lmargin=30mm",
                "rmargin=30mm",
                "top=25mm",
                "bmargin=25mm",
                "headheight=50mm",
            ],
        ),
        Package("microtype"),
        Package("authoraftertitle"),
        Package(" ", options=["MeX"]),
        # Package('geometry',options=['lmargin=25mm', 'rmargin=25mm',  'top=30mm', 'bmargin=25mm', 'headheight=50mm']),
        Package("listings"),
        Package("titlesec"),
        Package("fancyhdr"),
        Package("graphicx"),
        Package("indentfirst"),
        Package("pdfpages"),
        Package("amsmath"),
        Command("newcommand{\praca}", arguments=["Praca dyplomowa"]),
        Command("newcommand{\dyplom}", arguments=["Magisterska"]),
        Command("newcommand{\kierunek}", arguments=["Wpisać kierunek"]),
        Command("newcommand{\specjalnosc}", arguments=["Wpisać specjalność"]),
        Command("newcommand{\\autor}", arguments=["Imię i nazwisko autora"]),
        Command("newcommand{\opiekun}", arguments=["Wpisać opiekuna"]),
        Command("newcommand{\promotor}", arguments=["Wpisać promotora"]),
        Command("newcommand{\konsultant}", arguments=["Wpisać konsultanta"]),
        Command(
            "newcommand{\\tytul}", arguments=["Wpisać tytuł pracy dyplomowej po polsku"]
        ),
        Command("newcommand{\\album}", arguments=["Wpisać numer albumu"]),
        Command("newcommand{\supervisor}", arguments=["dr inż. Bogumił Chiliński"]),
        Command("newcommand{\\rok}", arguments=["Rok składania pracy"]),
        Command(
            "newcommand{\kluczowe}",
            arguments=["Słowa kluczowe: Wpisać słowa kluczowe po polsku"],
        ),
        # Command('renewcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']),
        Command("graphicspath{{../}}"),
        Command("frenchspacing"),
        Command("counterwithin{figure}{section}"),
        Command("counterwithin{table}{section}"),
        Command(
            "fancypagestyle{headings}{\\fancyhead{} \\renewcommand{\headrulewidth}{1pt} \\fancyheadoffset{0cm} \\fancyhead[RO]{\\nouppercase{\\leftmark}} \\fancyhead[LE]{\\nouppercase{\\leftmark}} \\fancyfoot{} \\fancyfoot[LE,RO]{\\thepage}}"
        ),
        Command(
            "fancypagestyle{plain}{\\fancyhf{} \\renewcommand{\\headrulewidth}{0pt} \\fancyfoot[LE,RO]{\\thepage}}"
        ),
        Command("numberwithin{equation}{section}"),
        Command("renewcommand{\\familydefault}{\\sfdefault}"),
        # \renewcommand{\familydefault}{\sfdefault}
    ]

    def __init__(
        self,
        default_filepath="default_filepath",
        title="Basic title",
        *,
        documentclass="article",
        document_options=["a4paper", "11pt", "twoside"],
        fontenc="T1",
        inputenc="utf8",
        font_size="normalsize",
        lmodern=False,
        textcomp=True,
        microtype=True,
        page_numbers=True,
        indent=None,
        geometry_options=[
            "inner=30mm",
            "outer=20mm",
            "bindingoffset=10mm",
            "top=25mm",
            "bottom=25mm",
        ],  # ,inner=20mm, outer=20mm, bindingoffset=10mm, top=25mm, bottom=25mm
        data=None,
    ):

        super().__init__(
            default_filepath=default_filepath,
            documentclass=documentclass,
            document_options=document_options,
            fontenc=fontenc,
            inputenc=inputenc,
            font_size=font_size,
            lmodern=lmodern,
            textcomp=textcomp,
            microtype=microtype,
            page_numbers=page_numbers,
            indent=indent,
            geometry_options=geometry_options,
            data=data,
        )
        #         label=self.label
        self.title = title
        # self.packages.append(Command('title', arguments=[NoEscape(self.title)]))
        #         self.packages.append(Command('date', arguments=[NoEscape('\\today')]))
        #         self.packages.append(Command('newcommand{\praca}', arguments=['Praca dyplomowa']))
        #         self.packages.append(Command('newcommand{\dyplom}', arguments=['Magisterska']))
        #         self.packages.append(Command('newcommand{\kierunek}', arguments=['Wpisać kierunek']))
        #         self.packages.append(Command('newcommand{\specjalnosc}', arguments=['Wpisać specjalność']))
        #         self.packages.append(Command('newcommand{\\autor}', arguments=['Imię i nazwisko autora']))
        #         self.packages.append(Command('newcommand{\opiekun}', arguments=['Wpisać opiekuna']))
        #         self.packages.append(Command('newcommand{\promotor}', arguments=['Wpisać promotora']))
        #         self.packages.append(Command('newcommand{\konsultant}', arguments=['Wpisać konsultanta']))
        #         self.packages.append(Command('newcommand{\\tytul}', arguments=['Wpisać tytuł pracy dyplomowej po polsku']))
        #         self.packages.append(Command('newcommand{\\album}', arguments=['303596']))
        #         self.packages.append(Command('newcommand{\supervisor}', arguments=['dr inż. Bogumił Chiliński']))
        #         self.packages.append(Command('newcommand{\\rok}', arguments=['Rok składania pracy']))
        #         self.packages.append(Command('newcommand{\kluczowe}', arguments=['Słowa kluczowe: Wpisać słowa kluczowe po polsku']))
        #         #self.packages.append(Command('renewcommand{\keywords}', arguments=['Keywords: Wpisać słowa kluczowe po angielsku']))
        self.packages.append(Command("graphicspath{{../}}"))
        #         self.append(Command('maketitle'))
        self.append(NoEscape("%%% New doc"))
        # tu implementować co tam potrzeba

    @classmethod
    def base_setup(cls, create_file=False):

        if create_file is True:
            return cls._create_base_setup_env()

        preliminary_str = """

Examplary setup is as follows:

#References to guide imports needed to include - to see the list of available guides insert and run the code below:
```{python}
from dynpy.utilities.creators import list_of_guides
list_of_guides()
```

## CELL 1
## Imports

```{python}
#Create file output
#Create file Images
#In file output create bibliography as .bib file (biblio.bib)

'''# File content begin


@misc{DynPi,
author={GitHub},
title="bogumilchilinski/dynpy",
url={https://github.com/bogumilchilinski/dynpy}",
note="Accessed:2024-05-04"
}


@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
'''# File content end

from dynpy.utilities.report import *
from dynpy.utilities.documents.document import StateOfArtReport


from dynpy.utilities.adaptable import TimeDataFrame
import pandas as pd


doc = StateOfArtReport('./output/report_name')
```



## CELL 2
## Document introduction

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_int = Section('Recenzja pracy dyplomowej pt.', numbering=False)
CurrentContainer(sec_int)

tytul= " Title of thesis"

display(ReportText(f' "{tytul} " '))
```
## CELL 3
## Content

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_1 = Section('Ocena zgodności treści pracy z jej tematem')
CurrentContainer(sec_1)

#czego_dotyczy = "rozbudowy biblioteki modeli dynamicznych w bibliotece DynPy"
czego_dotyczy = "weryfikacji empirycznej działania wahadłowego absorbera drgań"

#wklad_autora = "rozwój biblioteki systemów dynamicznych zawartych w DynPy, w tym implementację nowych klas, poprawę kodu istniejących przewodników po narzędziach oferowanych przez bibliotekę"
wklad_autora = "prowadzenie badań empirycznych dotyczących wahadłowych eliminatorów drgań"

#Jesli cos wiecej
Ponad_to = "Ponadto praca zawiera"
co_wiecej_zawiera = "funkcjonujący model 3D wykonany przez Autora pracy"

display(ReportText(f'''Rozważania zawarte w pracy dotyczącej {czego_dotyczy} odpowiadają jej tematowi. Praca składa się z siedmiu zasadniczych części. Pierwsza część jest poświęcona wprowadzeniu do problematyki, określeniu celu pracy, motywacji oraz metodyki badawczej. Druga część poświęcona jest stworzeniu modelu analitycznego układu drgającego, opisanego za pomocą równań ruchu wyprowadzonych metodą Lagrange’a. Model ten stanowił podstawę dalszych symulacji numerycznych. Trzeci rozdział dotyczy projektowania i realizacji stanowiska badawczego – od koncepcyjnego modelu w programie SOLIDWORKS po wykonanie rzeczywistej konstrukcji z systemem PTMD. Czwarta część pracy zawiera szczegółowy opis symulacji numerycznych przeprowadzonych w środowisku Python z użyciem biblioteki DynPy. Przeanalizowano wpływ takich parametrów jak długość wahadła i masa obciążenia na skuteczność tłumienia drgań. W rozdziale piątym zaprezentowano badania eksperymentalne wykonane na zaprojektowanym stanowisku. Szósta część zestawia i porównuje wyniki badań eksperymentalnych z rezultatami uzyskanymi w symulacjach, co pozwoliło na ocenę trafności modelu oraz skuteczności zastosowanego rozwiązania. Ostatnia część zawiera wnioski z przeprowadzonych prac i podsumowanie osiągniętych rezultatów. Opracowanie przedstawia wkład Autora w {wklad_autora}.'''))

## CELL 4
## Content

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_2 = Section('Ocena wartości merytorycznej, w szczególności wkładu własnego studenta oraz jasności jego przedstawienia')
CurrentContainer(sec_2)
#1 zdanie

#intro_way = "niejasny i nielogiczny "
intro_way = "jasny i logiczny"

#2 zdanie

#poziom_pracy = "wysokim"
poziom_pracy = 'odpowiednim'
#poziom_pracy = "niskim"

#zdanie
only_nec = "tylko niezbędne informacje, pozostając przy tym czytelne"

#radzenie_z_problemem = "Student dobrze radzi sobie z postawionymi problemami, potwierdza to przeprowadzona analiza systemu dynamicznego, przygotowanie raportu z analizy oraz implementacja nowych funkcjonalności."
radzenie_z_problemem = "Student dobrze radzi sobie z postawionymi problemami, potwierdza to przeprowadzona analiza i badanie systemu dynamicznego, przygotowanie oraz raportu z analizy."

#czego_to_wymagalo = "Zadania te wymagały wykazania się umiejętnością programowania oraz zarządzania strukturą kodu. Potwierdza to również umiejętności Autora w zakresie mechaniki, w szczególności dynamiki układów."
czego_to_wymagalo = "Zadania te wymagały wykazania się umiejętnością projektowania i realizacji badań empirycznych. Potwierdza to również kompetencje Autora w zakresie mechaniki, w szczególności dynamiki układów."

#lista uwag merytorycznych:

content_remarks = [
#"lapidarny wstęp, w którym brakuje zdefiniowania założeń, metodyki i motywacji",
#"brak wyszczególnienia wszystkich parametrów zastosowanych elementów wykonawczych",
#"brak obliczeń uzasadniających dobór wykorzystanych silników",
#"pominięcie obliczeń poboru mocy oraz doboru źródeł energii",
#"zastosowanie liniowej pojemności w modelu ogniwa elektrochemicznego",
"brak bezpośredniego porównania wyników z dowolnym rozwiązaniem referencyjnym (publikacje, inne oprogramownie inżynierskie)",
#"omówione poprawki w kodach mogłyby zostać przedstawione w szerszym kontekście",
"brak identyfikacji parametycznej modelu matematycznego oraz ostatecznie dobranych wartości parametrów symulacyjnych",
'widoczne róznice "wizualne" pomiędzy wynikami pomiarów oraz symulacji',
"lapidarne oraz nieprecyzyjne wyjaśnienie różnic w efektach procesu optymalizacji"
]

content_remarks_str = '\n \n - '+ ', \n \n - '.join(content_remarks) + '.'

display(ReportText(f'Autor przedstawia przeprowadzone prace w {intro_way} sposób. Analiza zawartości pracy pokazuje samodzielny wkład Studenta w postaci przeprowadzania procesu identyfikacji parametrycznej modelu dynamicznego. Wartość merytoryczna oraz sposób rozumowania są na dobrym poziomie. Opracowanie przekazuje czytelnikowi {only_nec}. {radzenie_z_problemem} {czego_to_wymagalo}  '))

display(Markdown(f'
Do głównych uwag merytorycznych należą:

{content_remarks_str}

'))

display(ReportText('Mimo to Autor dobrze radzi sobie z postawionymi problemami, potwierdzając to samodzielnym wykonaniem projektu.'))

## CELL 5
## Content

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_3 = Section('Ocena doboru i sposobu wykorzystania źródeł')
CurrentContainer(sec_3)
#liczba pozycji w literaturze
lit_numb = '13'
# w zależności od tego czego wymaga liczba wybieramy odpowiednia odmiane
pozycje = "pozycje"
pozycji = "pozycji"

#forma_zgodna = "Wszystkie pozycje bibliografii znajdują odwołanie w tekście. Forma odwołań jak i spis wykorzystanych publikacji oraz łączy nie są do końca zgodne z wymaganiami postawionymi w rozporządzeniu dot. pisania prac dyplomowych."

#forma_zgodna = "Wszystkie pozycje bibliografii znajdują odwołanie w tekście. Forma odwołań jak i spis wykorzystanych publikacji oraz łączy nie są do końca zgodne z wymaganiami postawionymi w rozporządzeniu dot. pisania prac dyplomowych."


# form_zgodna_puste_odwolania = "Niektóre odwołania w tekście nie znajdują odwołania w bibliografii. Forma odwołań jak i spis publikacji nie są zgodne z wytycznymi odnośnie formatowania źródeł literaturowych postawionymi w rozporządzeniu dot. pisania prac dyplomowych."

# brakuje_doi = "W przypadku niektórych publikacji, brakuje informacji o wydawcy, adresie doi lub dacie dostępu. "

forma_zgodna = "Wszystkie pozycje bibliografii znajdują odwołanie w tekście. Forma odwołań jak i spis wykorzystanych publikacji oraz łączy są zgodne z wymaganiami postawionymi w rozporządzeniu dot. pisania prac dyplomowych. "

display(ReportText(f'Na podstawie dokonanego przeglądu literaturowego można stwierdzić, że dobór pozycji nie budzi zastrzeżeń i jest adekwatny do poruszanej w pracy tematyki. Bibliografia liczy {lit_numb} {pozycji}. {forma_zgodna} '))


## CELL 6
## Content

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_4 = Section('Ocena trafności i spójności wniosków')
CurrentContainer(sec_4)

wnioski_log = "logiczny"
schemat_wnioskowania = "poprawny"
konkluzje = "zgodne z tematem pracy"
konk_a_cele= "odpowiadają postawionym celom"

lakoniczne = "Aczkolwiek należy zauważyć, że są one dość lakoniczne."
#lakoniczne = ''

display(ReportText(f'Autor pracy wnioskuje w sposób {wnioski_log}. Podsumowanie pracy zawiera streszczenie zrealizowanych zadań oraz wnioski wyciągnięte na ich podstawie. Schemat wnioskowania jest {schemat_wnioskowania}. Konkluzje są {konkluzje} i {konk_a_cele}. {lakoniczne}'))

## CELL 7
## Content

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_5 = Section('Ocena układu pracy (kolejność rozdziałów, sposób doboru materiału ilustracyjnego)')
CurrentContainer(sec_5)
l_rozdzial = "7"
l_stron = 49

display(ReportText(f'Praca składa się z {l_stron} stron, podzielonych na {l_rozdzial} rozdziałów i prowadzi czytelnika w sposób jasny i logiczny po omawianym temacie. Całość jest podzielona na część teoretyczną oraz praktyczną. Ostatni fragment pracy poświęcono podsumowaniu i wnioskom. Pozostałą część pracy stanowi bibliografia, spis rysunków oraz spis tabel. Dobór materiału ilustrującego nie budzi zastrzeżeń.'))

## CELL 8
## Remarks

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_6 = Section('Ocena redakcji pracy (poprawność języka, staranność opracowania edytorskiego)')

editorial_remarks = [
"sporadyczne stosowanie formy osobowej",

"brak wyszczególnienia wzorów i opisów zastosowanych symboli",

#"niepoprawne zastosowanie operatorów matematycznych np. '*' dla operacji mnożenia",

"brak spisu symboli na końcu pracy",

#"stosowanie różnego kroju pisma dla symboli matematycznych",

"stosowanie różnego kroju pisma dla symboli matematycznych lub jednostek",

#"występowanie grafik lub wypunktowania na końcu sekcji",

#"występowanie zwrotów nie mających wartości merytorycznej takich jak: oczywiście, niestety itp.",

"występowanie literówek",

"błędy stylistyczne i interpunkcyjne",

"miejscowy brak akpitów",

"brak odwołań do źródeł w przypadku niektórych grafik (zastosowano pośrednio w tekście)"

]


edit_remarks_str = '\n \n - '+ ', \n \n - '.join(editorial_remarks) + '.'

CurrentContainer(sec_6)

#poziom_pracy_form = "wysokim"
#poziom_pracy_form = 'odpowiednim'
poziom_pracy_form = 'zadowalającym'
#poziom_pracy_form = "niskim"

display(ReportText(f'Redakcja pracy jest na {poziom_pracy_form} poziomie. Praca utrzymana w bezosobowej formie i czasie przeszłym. Występujące odstępstwa od tej zasady nie wpływają na ocenę całego opracowania. Zastosowany język jest odpowiedni dla prezentowanego problemu.'))

display(Markdown(f'
Do głównych uwag edytorskich należą:

{edit_remarks_str}

'))



display(ReportText('Zastrzeżenia te nie wpływają istotnie na odbiór pracy oraz na jej wartość merytoryczną. Występujące błędy oraz wymienione uwagi zostały przekazane do informacji Autora.'))

# Występujące błędy stylistyczne i interpunkcyjne oraz wymienione uwagi nie wpływają na pozytywny odbiór pracy.

## CELL 9
## Remarks2

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_7 = Section('Ocena przestrzegania praw autorskich')
CurrentContainer(sec_7)
#1 zdanie
tworcza = "ma charakter twórczy"
nie_tworcza = "nie ma charakteru twórczego"

zastr = "budzi zastrzeżenia"
bez_zastr = "nie budzi zastrzeżeń"

#2 zdanie

byla_proba = "Zaobserwowano próbę"
nie_bylo_proby= "Nie zaobserwowano prób"

#3 zdanie

Anty_git = "Praca pomyślnie przeszła test antyplagiatowy w systemie JSA"
Anty_nope = "Praca nie przeszła testu antyplagiatowego w systemie JSA"

display(ReportText(f'Praca {tworcza} i samodzielność jej wykonania {bez_zastr}. {nie_bylo_proby} łamania praw autorskich. {Anty_git}. '))

## CELL 10
## Remarks3

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_8 = Section('Ocena umiejętności samodzielnego rozwiązywania zagadnień inżynierskich')
CurrentContainer(sec_8)

#display(ReportText(f'Dyplomant rozwiązał samodzielnie problem inżynierski. Wykazał się umiejętnością projektowania części przy użyciu modelowania geometrycznego oraz doboru komponentów mechatronicznych, a także opracowania oprogramowania inżynierskiego, potwierdzając to zawartym w pracy opisem zastosowanego oprogramowania, przeprowadzoną analizą wyników z zaprojektowanego urządzenia z wynikami walidacji poprawności działania urządzenia. Autor wykazał się umiejętnością kierowania i krytycznej oceny własnej pracy, co jest podstawą do rozwiązywania zagadnień inżynierskich.'))

outro = 'Dyplomant wykazał się również zdolnością do samodzielnego kierowania pracami oraz krytycznej oceny ich rezultatów, co stanowi kluczową podstawę do rozwiązywania problemów inżynierskich i dalszego rozwoju zawodowego w tej dziedzinie.'

display(ReportText(f'Dyplomant samodzielnie rozwiązał problem inżynierski, wykazując się zdolnościami analizy systemów dynamicznych oraz programowania. Zaprezentował swoje umiejętności podczas tworzenia i analizy symulacji systemów dynamicznych oraz prowadzenia badań empirycznych. Dyplomant zaprezentował oprogramowanie inżynierskie, które zostało opisane w pracy, a jego efektywność została zweryfikowana przez przeprowadzenie analiz i walidacji uzyskanych wyników. {outro}'))

## CELL 11
## Remarks4

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_9 = Section('Czy zakres i poziom pracy odpowiadają wymaganiom stawianym pracom dyplomowym odpowiedniego stopnia studiów?')
CurrentContainer(sec_9)
mgr= "magisterskim"
inz= " inżynierskim"
display(ReportText(f'Poziom pracy i jej zakres spełnia wymogi stawiane pracom {inz}. '))

## CELL 12
## Hints

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_10 = Section('Inne uwagi')
CurrentContainer(sec_10)
display(ReportText('Brak uwag. '))

## CELL 13
## Document

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

# Creating file
# Be sure *output* folder is in the current directory

Doc = Document(default_filepath='./output/doc_name')
Doc.append(sec_int)
Doc.append(sec_1)
Doc.append(sec_2)
Doc.append(sec_3)
Doc.append(sec_4)
Doc.append(sec_5)
Doc.append(sec_6)
Doc.append(sec_7)
Doc.append(sec_8)
Doc.append(sec_9)
Doc.append(sec_10)

Doc.generate_pdf(clean_tex = False))

# Generating file
doc.generate_pdf(clean_tex=True)
```

"""

        display(IPMarkdown(preliminary_str))

        preliminary_str = """

#Example:

#To prepare a simple document with text and images:

#Good practice here is to allocate 1 section per 1 cell



## CELL 1
## Imports

```{python}
#Create file output
#Create file Images
#In file output create bibliography as .bib file (biblio.bib)

'''# File content begin


@misc{DynPi,
author={GitHub},
title="bogumilchilinski/dynpy",
url={https://github.com/bogumilchilinski/dynpy}",
note="Accessed:2024-05-04"
}


@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
'''# File content end

from dynpy.utilities.report import *
from dynpy.utilities.documents.document import StateOfArtReport


from dynpy.utilities.adaptable import TimeDataFrame
import pandas as pd


doc = StateOfArtReport('./output/report_name')
```



## CELL 2
## Document introduction

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_int = Section('Recenzja pracy dyplomowej pt.', numbering=False)
CurrentContainer(sec_int)

tytul= " Title of thesis"

display(ReportText(f' "{tytul} " '))
```
## CELL 3
## Content

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_1 = Section('Ocena zgodności treści pracy z jej tematem')
CurrentContainer(sec_1)

#czego_dotyczy = "rozbudowy biblioteki modeli dynamicznych w bibliotece DynPy"
czego_dotyczy = "weryfikacji empirycznej działania wahadłowego absorbera drgań"

#wklad_autora = "rozwój biblioteki systemów dynamicznych zawartych w DynPy, w tym implementację nowych klas, poprawę kodu istniejących przewodników po narzędziach oferowanych przez bibliotekę"
wklad_autora = "prowadzenie badań empirycznych dotyczących wahadłowych eliminatorów drgań"

#Jesli cos wiecej
Ponad_to = "Ponadto praca zawiera"
co_wiecej_zawiera = "funkcjonujący model 3D wykonany przez Autora pracy"

display(ReportText(f'''Rozważania zawarte w pracy dotyczącej {czego_dotyczy} odpowiadają jej tematowi. Praca składa się z siedmiu zasadniczych części. Pierwsza część jest poświęcona wprowadzeniu do problematyki, określeniu celu pracy, motywacji oraz metodyki badawczej. Druga część poświęcona jest stworzeniu modelu analitycznego układu drgającego, opisanego za pomocą równań ruchu wyprowadzonych metodą Lagrange’a. Model ten stanowił podstawę dalszych symulacji numerycznych. Trzeci rozdział dotyczy projektowania i realizacji stanowiska badawczego – od koncepcyjnego modelu w programie SOLIDWORKS po wykonanie rzeczywistej konstrukcji z systemem PTMD. Czwarta część pracy zawiera szczegółowy opis symulacji numerycznych przeprowadzonych w środowisku Python z użyciem biblioteki DynPy. Przeanalizowano wpływ takich parametrów jak długość wahadła i masa obciążenia na skuteczność tłumienia drgań. W rozdziale piątym zaprezentowano badania eksperymentalne wykonane na zaprojektowanym stanowisku. Szósta część zestawia i porównuje wyniki badań eksperymentalnych z rezultatami uzyskanymi w symulacjach, co pozwoliło na ocenę trafności modelu oraz skuteczności zastosowanego rozwiązania. Ostatnia część zawiera wnioski z przeprowadzonych prac i podsumowanie osiągniętych rezultatów. Opracowanie przedstawia wkład Autora w {wklad_autora}.'''))

## CELL 4
## Content

#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_2 = Section('Ocena wartości merytorycznej, w szczególności wkładu własnego studenta oraz jasności jego przedstawienia')
CurrentContainer(sec_2)
#1 zdanie

#intro_way = "niejasny i nielogiczny "
intro_way = "jasny i logiczny"

#2 zdanie

#poziom_pracy = "wysokim"
poziom_pracy = 'odpowiednim'
#poziom_pracy = "niskim"

#zdanie
only_nec = "tylko niezbędne informacje, pozostając przy tym czytelne"

#radzenie_z_problemem = "Student dobrze radzi sobie z postawionymi problemami, potwierdza to przeprowadzona analiza systemu dynamicznego, przygotowanie raportu z analizy oraz implementacja nowych funkcjonalności."
radzenie_z_problemem = "Student dobrze radzi sobie z postawionymi problemami, potwierdza to przeprowadzona analiza i badanie systemu dynamicznego, przygotowanie oraz raportu z analizy."

#czego_to_wymagalo = "Zadania te wymagały wykazania się umiejętnością programowania oraz zarządzania strukturą kodu. Potwierdza to również umiejętności Autora w zakresie mechaniki, w szczególności dynamiki układów."
czego_to_wymagalo = "Zadania te wymagały wykazania się umiejętnością projektowania i realizacji badań empirycznych. Potwierdza to również kompetencje Autora w zakresie mechaniki, w szczególności dynamiki układów."

#lista uwag merytorycznych:

content_remarks = [
#"lapidarny wstęp, w którym brakuje zdefiniowania założeń, metodyki i motywacji",
#"brak wyszczególnienia wszystkich parametrów zastosowanych elementów wykonawczych",
#"brak obliczeń uzasadniających dobór wykorzystanych silników",
#"pominięcie obliczeń poboru mocy oraz doboru źródeł energii",
#"zastosowanie liniowej pojemności w modelu ogniwa elektrochemicznego",
"brak bezpośredniego porównania wyników z dowolnym rozwiązaniem referencyjnym (publikacje, inne oprogramownie inżynierskie)",
#"omówione poprawki w kodach mogłyby zostać przedstawione w szerszym kontekście",
"brak identyfikacji parametycznej modelu matematycznego oraz ostatecznie dobranych wartości parametrów symulacyjnych",
'widoczne róznice "wizualne" pomiędzy wynikami pomiarów oraz symulacji',
"lapidarne oraz nieprecyzyjne wyjaśnienie różnic w efektach procesu optymalizacji"
]

content_remarks_str = '\n \n - '+ ', \n \n - '.join(content_remarks) + '.'

display(ReportText(f'Autor przedstawia przeprowadzone prace w {intro_way} sposób. Analiza zawartości pracy pokazuje samodzielny wkład Studenta w postaci przeprowadzania procesu identyfikacji parametrycznej modelu dynamicznego. Wartość merytoryczna oraz sposób rozumowania są na dobrym poziomie. Opracowanie przekazuje czytelnikowi {only_nec}. {radzenie_z_problemem} {czego_to_wymagalo}  '))

display(Markdown(f'
Do głównych uwag merytorycznych należą:

{content_remarks_str}

'))

display(ReportText('Mimo to Autor dobrze radzi sobie z postawionymi problemami, potwierdzając to samodzielnym wykonaniem projektu.'))

## CELL 5
## Content

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_3 = Section('Ocena doboru i sposobu wykorzystania źródeł')
CurrentContainer(sec_3)
#liczba pozycji w literaturze
lit_numb = '13'
# w zależności od tego czego wymaga liczba wybieramy odpowiednia odmiane
pozycje = "pozycje"
pozycji = "pozycji"

#forma_zgodna = "Wszystkie pozycje bibliografii znajdują odwołanie w tekście. Forma odwołań jak i spis wykorzystanych publikacji oraz łączy nie są do końca zgodne z wymaganiami postawionymi w rozporządzeniu dot. pisania prac dyplomowych."

#forma_zgodna = "Wszystkie pozycje bibliografii znajdują odwołanie w tekście. Forma odwołań jak i spis wykorzystanych publikacji oraz łączy nie są do końca zgodne z wymaganiami postawionymi w rozporządzeniu dot. pisania prac dyplomowych."


# form_zgodna_puste_odwolania = "Niektóre odwołania w tekście nie znajdują odwołania w bibliografii. Forma odwołań jak i spis publikacji nie są zgodne z wytycznymi odnośnie formatowania źródeł literaturowych postawionymi w rozporządzeniu dot. pisania prac dyplomowych."

# brakuje_doi = "W przypadku niektórych publikacji, brakuje informacji o wydawcy, adresie doi lub dacie dostępu. "

forma_zgodna = "Wszystkie pozycje bibliografii znajdują odwołanie w tekście. Forma odwołań jak i spis wykorzystanych publikacji oraz łączy są zgodne z wymaganiami postawionymi w rozporządzeniu dot. pisania prac dyplomowych. "

display(ReportText(f'Na podstawie dokonanego przeglądu literaturowego można stwierdzić, że dobór pozycji nie budzi zastrzeżeń i jest adekwatny do poruszanej w pracy tematyki. Bibliografia liczy {lit_numb} {pozycji}. {forma_zgodna} '))


## CELL 6
## Content

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_4 = Section('Ocena trafności i spójności wniosków')
CurrentContainer(sec_4)

wnioski_log = "logiczny"
schemat_wnioskowania = "poprawny"
konkluzje = "zgodne z tematem pracy"
konk_a_cele= "odpowiadają postawionym celom"

lakoniczne = "Aczkolwiek należy zauważyć, że są one dość lakoniczne."
#lakoniczne = ''

display(ReportText(f'Autor pracy wnioskuje w sposób {wnioski_log}. Podsumowanie pracy zawiera streszczenie zrealizowanych zadań oraz wnioski wyciągnięte na ich podstawie. Schemat wnioskowania jest {schemat_wnioskowania}. Konkluzje są {konkluzje} i {konk_a_cele}. {lakoniczne}'))

## CELL 7
## Content

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_5 = Section('Ocena układu pracy (kolejność rozdziałów, sposób doboru materiału ilustracyjnego)')
CurrentContainer(sec_5)
l_rozdzial = "7"
l_stron = 49

display(ReportText(f'Praca składa się z {l_stron} stron, podzielonych na {l_rozdzial} rozdziałów i prowadzi czytelnika w sposób jasny i logiczny po omawianym temacie. Całość jest podzielona na część teoretyczną oraz praktyczną. Ostatni fragment pracy poświęcono podsumowaniu i wnioskom. Pozostałą część pracy stanowi bibliografia, spis rysunków oraz spis tabel. Dobór materiału ilustrującego nie budzi zastrzeżeń.'))

## CELL 8
## Remarks

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_6 = Section('Ocena redakcji pracy (poprawność języka, staranność opracowania edytorskiego)')

editorial_remarks = [
"sporadyczne stosowanie formy osobowej",

"brak wyszczególnienia wzorów i opisów zastosowanych symboli",

#"niepoprawne zastosowanie operatorów matematycznych np. '*' dla operacji mnożenia",

"brak spisu symboli na końcu pracy",

#"stosowanie różnego kroju pisma dla symboli matematycznych",

"stosowanie różnego kroju pisma dla symboli matematycznych lub jednostek",

#"występowanie grafik lub wypunktowania na końcu sekcji",

#"występowanie zwrotów nie mających wartości merytorycznej takich jak: oczywiście, niestety itp.",

"występowanie literówek",

"błędy stylistyczne i interpunkcyjne",

"miejscowy brak akpitów",

"brak odwołań do źródeł w przypadku niektórych grafik (zastosowano pośrednio w tekście)"

]


edit_remarks_str = '\n \n - '+ ', \n \n - '.join(editorial_remarks) + '.'

CurrentContainer(sec_6)

#poziom_pracy_form = "wysokim"
#poziom_pracy_form = 'odpowiednim'
poziom_pracy_form = 'zadowalającym'
#poziom_pracy_form = "niskim"

display(ReportText(f'Redakcja pracy jest na {poziom_pracy_form} poziomie. Praca utrzymana w bezosobowej formie i czasie przeszłym. Występujące odstępstwa od tej zasady nie wpływają na ocenę całego opracowania. Zastosowany język jest odpowiedni dla prezentowanego problemu.'))

display(Markdown(f'
Do głównych uwag edytorskich należą:

{edit_remarks_str}

'))



display(ReportText('Zastrzeżenia te nie wpływają istotnie na odbiór pracy oraz na jej wartość merytoryczną. Występujące błędy oraz wymienione uwagi zostały przekazane do informacji Autora.'))

# Występujące błędy stylistyczne i interpunkcyjne oraz wymienione uwagi nie wpływają na pozytywny odbiór pracy.

## CELL 9
## Remarks2

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_7 = Section('Ocena przestrzegania praw autorskich')
CurrentContainer(sec_7)
#1 zdanie
tworcza = "ma charakter twórczy"
nie_tworcza = "nie ma charakteru twórczego"

zastr = "budzi zastrzeżenia"
bez_zastr = "nie budzi zastrzeżeń"

#2 zdanie

byla_proba = "Zaobserwowano próbę"
nie_bylo_proby= "Nie zaobserwowano prób"

#3 zdanie

Anty_git = "Praca pomyślnie przeszła test antyplagiatowy w systemie JSA"
Anty_nope = "Praca nie przeszła testu antyplagiatowego w systemie JSA"

display(ReportText(f'Praca {tworcza} i samodzielność jej wykonania {bez_zastr}. {nie_bylo_proby} łamania praw autorskich. {Anty_git}. '))

## CELL 10
## Remarks3

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_8 = Section('Ocena umiejętności samodzielnego rozwiązywania zagadnień inżynierskich')
CurrentContainer(sec_8)

#display(ReportText(f'Dyplomant rozwiązał samodzielnie problem inżynierski. Wykazał się umiejętnością projektowania części przy użyciu modelowania geometrycznego oraz doboru komponentów mechatronicznych, a także opracowania oprogramowania inżynierskiego, potwierdzając to zawartym w pracy opisem zastosowanego oprogramowania, przeprowadzoną analizą wyników z zaprojektowanego urządzenia z wynikami walidacji poprawności działania urządzenia. Autor wykazał się umiejętnością kierowania i krytycznej oceny własnej pracy, co jest podstawą do rozwiązywania zagadnień inżynierskich.'))

outro = 'Dyplomant wykazał się również zdolnością do samodzielnego kierowania pracami oraz krytycznej oceny ich rezultatów, co stanowi kluczową podstawę do rozwiązywania problemów inżynierskich i dalszego rozwoju zawodowego w tej dziedzinie.'

display(ReportText(f'Dyplomant samodzielnie rozwiązał problem inżynierski, wykazując się zdolnościami analizy systemów dynamicznych oraz programowania. Zaprezentował swoje umiejętności podczas tworzenia i analizy symulacji systemów dynamicznych oraz prowadzenia badań empirycznych. Dyplomant zaprezentował oprogramowanie inżynierskie, które zostało opisane w pracy, a jego efektywność została zweryfikowana przez przeprowadzenie analiz i walidacji uzyskanych wyników. {outro}'))

## CELL 11
## Remarks4

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_9 = Section('Czy zakres i poziom pracy odpowiadają wymaganiom stawianym pracom dyplomowym odpowiedniego stopnia studiów?')
CurrentContainer(sec_9)
mgr= "magisterskim"
inz= " inżynierskim"
display(ReportText(f'Poziom pracy i jej zakres spełnia wymogi stawiane pracom {inz}. '))

## CELL 12
## Hints

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

sec_10 = Section('Inne uwagi')
CurrentContainer(sec_10)
display(ReportText('Brak uwag. '))

## CELL 13
## Document

```{python}
#!!! BE SURE ALL PREVIOUS CELLS ARE RUN !!!#
#!!!       BECAUSE OF NEEDED IMPORTS    !!!#

# Creating file
# Be sure *output* folder is in the current directory

Doc = Document(default_filepath='./output/doc_name')
Doc.append(sec_int)
Doc.append(sec_1)
Doc.append(sec_2)
Doc.append(sec_3)
Doc.append(sec_4)
Doc.append(sec_5)
Doc.append(sec_6)
Doc.append(sec_7)
Doc.append(sec_8)
Doc.append(sec_9)
Doc.append(sec_10)

Doc.generate_pdf(clean_tex = False))

# Generating file
doc.generate_pdf(clean_tex=True)
```


"""
        return ObjectCode(preliminary_str)

    Jupyter_file_content = """{
    "cells": [
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "d3000c",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "from dynpy.utilities.report import *",
                "\\n",
                "from dynpy.utilities.templates.document import WutThesis",
                "\\n",
                "",
                "doc = WutThesis('./output/thesis_name')"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "c3d54b",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "doc.base_setup()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "11531d",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": []
        }
    ],
    "metadata": {
        "kernelspec": {
            "argv": [
                "/usr/bin/python3",
                "-m",
                "ipykernel",
                "--HistoryManager.enabled=False",
                "--matplotlib=inline",
                "-c",
                "%config InlineBackend.figure_formats = set(['retina'])import matplotlib; matplotlib.rcParams['figure.figsize'] = (12, 7)",
                "-f",
                "{connection_file}"
            ],
            "display_name": "Python 3 (system-wide)",
            "env": {},
            "language": "python",
            "metadata": {
                "cocalc": {
                    "description": "Python 3 programming language",
                    "priority": 100,
                    "url": "https://www.python.org/"
                }
            },
            "name": "python3",
            "resource_dir": "/ext/jupyter/kernels/python3"
        },
        "language_info": {
            "codemirror_mode": {
                "name": "ipython",
                "version": 3
            },
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "nbconvert_exporter": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}"""
    articles_bib = """@INPROCEEDINGS{brief_history_of_simulations,
  author={Goldsman, David and Nance, Richard E. and Wilson, James R.},
  booktitle={Proceedings of the 2010 Winter Simulation Conference},
  title={A brief history of simulation revisited},
  year={2010},
  volume={},
  number={},
  pages={567-574},
  doi={10.1109/WSC.2010.5679129}
}
@article{eckhardt1987stan,
  title={Stan Ulam, John von Neumann},
  author={Eckhardt, Roger},
  journal={Los Alamos Science},
  volume={100},
  number={15},
  pages={131},
  year={1987},
  publisher={Los Alamos Scientific Laboratory}
}
@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
@misc{RTPW,
  author = {mgr Tomasz Duda},
  title = {Rysunek Techniczny Podstawowe Wiadomości},
  journal = {},
  year = {},
  number = {},
  pages = {10},
  doi = {}
}
"""

    def _create_directory(path):
        """
        Method create directory from given path

        Arguments:
        path - (str), name of directory that method will create, if subdirectories shall be created, path should look like this: 'directory/subdirectory'
        """
        import os

        try:
            os.makedirs(path)
            print(f"Directory '{path}' created")

        except FileExistsError:
            print(f"Directory '{path}' exists.")

        except PermissionError:
            print(f"Permission denied, you don't have sudo permission")

        except Exception as e:
            print(f"An error occurred: {e}")

    def _create_file(name, content=None, path=None):
        """
        Method create file and write content to it

        Arguments:
        name - (str), name of file created by method,
        content - (str), optional,  string with a content of file, that is going to be writen to it,
        path - (str), optional, directory where file should be created
        """
        import os

        if path is not None:
            if not os.path.exists(path):
                _create_directory(path)

            with open(os.path.join(path, name), "w") as file:
                file.write(content)

        else:
            with open(name, "w") as file:
                file.write(content)

        file.close()

    @classmethod
    def _create_base_setup_env(cls):
        """
        Method that create Jupyter notebook file with WUT thesis base setup and output directory with .bib file
        """

        Jupyter_file_content = """{
    "cells": [
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "d3000c",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "from dynpy.utilities.report import *",
                "\\n",
                "from dynpy.utilities.templates.document import WutThesis",
                "\\n",
                "",
                "doc = WutThesis('./output/thesis_name')"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "c3d54b",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": [
                "doc.base_setup()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": 0,
            "id": "11531d",
            "metadata": {
                "collapsed": false
            },
            "outputs": [],
            "source": []
        }
    ],
    "metadata": {
        "kernelspec": {
            "argv": [
                "/usr/bin/python3",
                "-m",
                "ipykernel",
                "--HistoryManager.enabled=False",
                "--matplotlib=inline",
                "-c",
                "%config InlineBackend.figure_formats = set(['retina'])import matplotlib; matplotlib.rcParams['figure.figsize'] = (12, 7)",
                "-f",
                "{connection_file}"
            ],
            "display_name": "Python 3 (system-wide)",
            "env": {},
            "language": "python",
            "metadata": {
                "cocalc": {
                    "description": "Python 3 programming language",
                    "priority": 100,
                    "url": "https://www.python.org/"
                }
            },
            "name": "python3",
            "resource_dir": "/ext/jupyter/kernels/python3"
        },
        "language_info": {
            "codemirror_mode": {
                "name": "ipython",
                "version": 3
            },
            "file_extension": ".py",
            "mimetype": "text/x-python",
            "name": "python",
            "nbconvert_exporter": "python",
            "pygments_lexer": "ipython3",
            "version": "3.10.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 4
}"""
        articles_bib = """@INPROCEEDINGS{brief_history_of_simulations,
  author={Goldsman, David and Nance, Richard E. and Wilson, James R.},
  booktitle={Proceedings of the 2010 Winter Simulation Conference},
  title={A brief history of simulation revisited},
  year={2010},
  volume={},
  number={},
  pages={567-574},
  doi={10.1109/WSC.2010.5679129}
}
@article{eckhardt1987stan,
  title={Stan Ulam, John von Neumann},
  author={Eckhardt, Roger},
  journal={Los Alamos Science},
  volume={100},
  number={15},
  pages={131},
  year={1987},
  publisher={Los Alamos Scientific Laboratory}
}
@book{lutz2001programming,
  title={Programming python},
  author={Lutz, Mark},
  year={2001},
  publisher={" O'Reilly Media, Inc."}
}
@misc{NumPy, url={https://numpy.org/}, journal={NumPy}}
@misc{pandas, url={https://pandas.pydata.org/}, journal={pandas}}
@misc{RTPW,
  author = {mgr Tomasz Duda},
  title = {Rysunek Techniczny Podstawowe Wiadomości},
  journal = {},
  year = {},
  number = {},
  pages = {10},
  doi = {}
}
"""
        cls._create_file("WUT_Thesis_starter.ipynb", Jupyter_file_content)

        os.makedirs("output")
        os.makedirs("tikzplots")

        cls._create_file("articles.bib", articles_bib, "output")






# --- ThesisSheet Generator --- #
from pathlib import Path
from datetime import datetime
import shutil, subprocess




class ThesisCardMini:

    def __init__(self, template_path: Path | None = None, output_dir: Path | None = None, try_pdf: bool = True):
        if template_path is None:
            template_path = Path(__file__).resolve().parent / "documents" / "thesisCard.docx"
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Nie znaleziono szablonu: {self.template_path}")

        self.output_dir = Path(output_dir or (Path.cwd() / "build"))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.try_pdf = try_pdf

    @staticmethod
    def _canon(s: str) -> str:
        return (s or "").replace("\u00a0", " ").strip().lower().rstrip(":")

    def _set_label_in_paragraphs(self, doc, label: str, value: str) -> bool:

        lab = self._canon(label)
        for p in doc.paragraphs:
            if self._canon(p.text).startswith(lab):
                before = p.text.split(":", 1)[0] if ":" in p.text else label
                p.text = f"{before.strip()}: {value.strip()}"
                return True
        return False

    def _set_label_in_tables(self, doc, label: str, value: str) -> bool:

        lab = self._canon(label)
        for table in doc.tables:
            for row in table.rows:
                for j, cell in enumerate(row.cells):
                    txt = cell.text or ""
                    if self._canon(txt).startswith(lab):
                        if j + 1 < len(row.cells):
                            row.cells[j + 1].text = value.strip()
                        else:
                            before = txt.split(":", 1)[0] if ":" in txt else label
                            cell.text = f"{before.strip()}: {value.strip()}"
                        return True
        return False

    def _set_value_below_label(self, doc, label: str, value: str) -> bool:

        lab = self._canon(label)

        # akapity
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            if self._canon(p.text).startswith(lab):
                if i + 1 < len(paras):
                    paras[i + 1].text = value.strip()
                    return True
                return False

        # tabele
        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if self._canon(cell.text).startswith(lab):
                        rr = r_idx + 1
                        if rr < len(table.rows):
                            table.rows[rr].cells[c_idx].text = value.strip()
                            return True
                        if c_idx + 1 < len(row.cells):
                            row.cells[c_idx + 1].text = value.strip()
                            return True
                        before = (cell.text or label).split(":", 1)[0]
                        cell.text = f"{before.strip()}:\n{value.strip()}"
                        return True
        return False

    def _set_below_any_label(self, doc, labels: list[str], value: str) -> bool:

        for lab in labels:
            if self._set_value_below_label(doc, lab, value):
                return True
        return False

    def _set_tasks_after_label(self, doc, label: str, tasks: list[str]) -> bool:

        lab = self._canon(label)

        # akapity
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            if self._canon(p.text).startswith(lab):
                for k in range(5):
                    idx = i + 1 + k
                    if idx < len(paras):
                        txt = tasks[k].strip() if k < len(tasks) else ""
                        paras[idx].text = f"{k+1}. {txt}".rstrip()
                return True

        # tabele
        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if self._canon(cell.text).startswith(lab):
                        for k in range(5):
                            rr = r_idx + 1 + k
                            if rr < len(table.rows):
                                txt = tasks[k].strip() if k < len(tasks) else ""
                                table.rows[rr].cells[c_idx].text = f"{k+1}. {txt}".rstrip()
                        return True
        return False

    def _insert_date_above_label(self, doc, label: str | list[str], date_text: str, align: str = "center") -> bool:


        try:
            from docx import Document
        except Exception as e:
            raise RuntimeError("Zainstaluj zależność: pip install python-docx") from e

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        labels = [label] if isinstance(label, str) else label
        labs = [self._canon(l) for l in labels]

        def _align_para(p, how: str):
            if how == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif how == "left":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        for p in doc.paragraphs:
            if any(l in self._canon(p.text) for l in labs):
                try:
                    new_p = p.insert_paragraph_before(date_text)
                except Exception:
                    p.text = f"{date_text}\n{p.text}"
                    _align_para(p, align)
                    return True
                _align_para(new_p, align)
                return True


        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if any(l in self._canon(cell.text) for l in labs):
                        # wiersz powyżej → ta sama kolumna
                        if r_idx > 0:
                            above = table.rows[r_idx - 1].cells[c_idx]
                            above.text = date_text
                            for para in above.paragraphs:
                                _align_para(para, align)
                            return True
                        # fallback: brak wiersza powyżej → prepend w tej samej komórce
                        if cell.text.strip():
                            cell.text = f"{date_text}\n{cell.text}"
                        else:
                            cell.text = date_text
                        for para in cell.paragraphs:
                            _align_para(para, align)
                        return True
        return False



    def _convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        # 1) docx2pdf
        try:
            from docx2pdf import convert as docx2pdf_convert  # type: ignore
            docx2pdf_convert(str(docx_path), str(pdf_path))
            return pdf_path.exists()
        except Exception:
            pass
        # 2) LibreOffice
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            try:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                return pdf_path.exists()
            except Exception:
                return False
        return False



    def make_pdf(
        self,
        rodzaj: str = "niestacjonarne",
        stopien: str = "II",
        kierunek: str | None = None,
        specjalnosc: str | None = None,
        dyplomant: str | None = None,
        numer_albumu: str | None = None,
        email: str | None = None,
        prowadzacy: str | None = None,
        konsultant: str | None = None,
        temat: str | None = None,              # Temat POD etykietą (nowa linia)
        opis: str | None = None,               # Opis POD etykietą (nowa linia)
        jezyk: str | None = "Polski",          # Język opracowania po dwukropku
        zadania: list[str] | None = None,      # 1..5 zadań, opcjonalnie
        wstaw_date: bool = True,
        data_aktualna: str | None = None,      # np. "10.10.2025"; None -> dziś
        date_align: str = "center",
        signPath: str | Path | None = None,
        sign_x_mm: float | None = None,   # X od lewego-górnego rogu
        sign_y_mm: float | None = None,   # Y od lewego-górnego rogu
        sign_w_mm: float = 45.0,          # szerokość obrazka
        sign_h_mm: float | None = None,   # None = zachowaj proporcje                        
                                
                                                        # "center" (domyślnie) / "right" / "left"
    ) -> dict:
        from docx import Document
        doc = Document(str(self.template_path))

        # 1) standardowe pola 'po dwukropku'/ jeżeli zmieni się docelowy plik karty pracy, to może być konieczne dostosowanie tej części - na razie zostawiam jak jest.
        fields = {
            "Rodzaj studiów": rodzaj,
            "Stopień studiów": stopien,
            "Kierunek": kierunek,
            "Specjalność": specjalnosc,
            "Dyplomant": dyplomant,
            "Numer albumu": numer_albumu,
            "Adres e-mail": email,
            "Prowadzący": prowadzacy,
            "Konsultant": konsultant,
            "Język opracowania": jezyk,
        }

        not_found: list[str] = []
        for label, value in fields.items():
            if value is None or (isinstance(value, str) and value.strip() == ""):
                continue
            ok = self._set_label_in_paragraphs(doc, label, value)
            if not ok:
                ok = self._set_label_in_tables(doc, label, value)
            if not ok:
                not_found.append(label)


        if isinstance(temat, str) and temat.strip():
            ok_topic = self._set_below_any_label(
                doc,
                ["Temat pracy dyplomowej (nowy / zmieniony)*", "Temat pracy dyplomowej", "Temat pracy"],
                temat.strip(),
            )
            if not ok_topic:
                not_found.append("Temat pracy")


        if isinstance(opis, str) and opis.strip():
            opis_one_line = opis.replace("\r", " ").replace("\n", " ").strip()
            ok_opis = self._set_below_any_label(
                doc,
                ["Zwięzły opis celu pracy", "Zwięzły opis pracy", "Opis pracy"],
                opis_one_line,
            )
            if not ok_opis:
                not_found.append("Zwięzły opis celu pracy")


        if zadania:
            tasks = (zadania or [])[:5]
            ok_tasks = self._set_tasks_after_label(doc, "Główne zadania do wykonania", tasks)
            if not ok_tasks:
                not_found.append("Główne zadania do wykonania")

        if wstaw_date:
            today_raw = data_aktualna or datetime.now().strftime("%d.%m.%Y")
            trail = "\u00A0" * 37
            today = f"{today_raw}{trail}"

            ok_date = self._insert_date_above_label(
                doc,
                ["Data wydania", "Data wydania:", "data wydania"],
                today,
                align=date_align,
            )
            if not ok_date:
                not_found.append("Data wydania (wstawienie daty nad)")

        if not_found:
            raise ValueError("Nie znaleziono etykiet w szablonie: " + ", ".join(not_found))

        
        out_docx = self.output_dir / f"ThesisCard.docx"
        out_pdf  = self.output_dir / f"ThesisCard.pdf"

        doc.save(str(out_docx))
        made_pdf = self.try_pdf and self._convert_to_pdf(out_docx, out_pdf)

        # Add picture if path and picture are given
        if made_pdf and signPath:
            try:
                self.signPdf(
                    pdf_path=out_pdf,
                    sign_path=Path(signPath),
                    x_mm=sign_x_mm, y_mm=sign_y_mm,
                    width_mm=sign_w_mm, height_mm=sign_h_mm,
                )
            except Exception as e:
                return {"docx": str(out_docx), "pdf": str(out_pdf), "signature_error": str(e)}
    



    def signPdf(
        self,
        pdf_path: Path,
        sign_path: Path,
        *,
        x_mm: float | None,          
        y_mm: float | None,          
        width_mm: float = 45.0,
        height_mm: float | None = None,  # None => proporcje z obrazka
    ) -> bool:

        from pathlib import Path
        import io, os, time
        pdf_path = Path(pdf_path); sign_path = Path(sign_path)
        if not pdf_path.exists(): raise FileNotFoundError(f"PDF nie istnieje: {pdf_path}")
        if not sign_path.exists(): raise FileNotFoundError(f"Obraz podpisu nie istnieje: {sign_path}")
        if x_mm is None or y_mm is None:
            raise ValueError("Podaj x_mm i y_mm (mm od lewego-górnego rogu).")

        def mm2pt(mm: float) -> float: return mm * 72.0 / 25.4

        # --- Preferowany wariant: PyMuPDF (overlay + saveIncr) ---
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(pdf_path))
            page = doc[0]
            # rozmiar obrazka (do zachowania proporcji)
            try:
                pm = fitz.Pixmap(str(sign_path)); iw, ih = pm.width, pm.height; pm = None
            except Exception:
                iw, ih = 1, 1

            w_pt = mm2pt(width_mm)
            h_pt = mm2pt(height_mm) if height_mm else w_pt * (ih / iw)

            x0 = mm2pt(x_mm)
            y0 = mm2pt(y_mm)
            rect = fitz.Rect(x0, y0, x0 + w_pt, y0 + h_pt)

            page.insert_image(rect, filename=str(sign_path), keep_proportion=True, overlay=True)

            try:
                doc.saveIncr(); doc.close()
            except Exception:
                tmp = pdf_path.with_suffix(".signed.tmp.pdf")
                doc.save(str(tmp)); doc.close()
                try:
                    os.replace(tmp, pdf_path)
                finally:
                    if tmp.exists(): tmp.unlink(missing_ok=True)
            return True
        except Exception:
            # --- Fallback: ReportLab + PyPDF2 (też overlay) ---
            try:
                from PyPDF2 import PdfReader, PdfWriter
                from reportlab.pdfgen import canvas
                from reportlab.lib.utils import ImageReader
                import io

                reader = PdfReader(str(pdf_path))
                page0 = reader.pages[0]
                pw, ph = float(page0.mediabox.width), float(page0.mediabox.height)

                img = ImageReader(str(sign_path))
                iw, ih = img.getSize()
                w_pt = mm2pt(width_mm)
                h_pt = mm2pt(height_mm) if height_mm else w_pt * (ih / iw)

                # ReportLab ma (0,0) w lewym-DOLNYM rogu => konwersja Y:
                x = mm2pt(x_mm)
                y = ph - mm2pt(y_mm) - h_pt

                buf = io.BytesIO()
                c = canvas.Canvas(buf, pagesize=(pw, ph))
                c.drawImage(img, x, y, width=w_pt, height=h_pt, preserveAspectRatio=True, mask='auto')
                c.save(); buf.seek(0)

                overlay = PdfReader(buf).pages[0]
                page0.merge_page(overlay)

                writer = PdfWriter()
                for i, p in enumerate(reader.pages):
                    writer.add_page(page0 if i == 0 else p)

                tmp = pdf_path.with_suffix(".signed.tmp.pdf")
                with open(tmp, "wb") as f: writer.write(f)
                try:
                    os.replace(tmp, pdf_path)
                finally:
                    if os.path.exists(tmp): os.remove(tmp)
                return True
            except Exception:
                return False