# from .guides import (UsageOfDynamicSystemsGuide, Guide, EngeneeringDrawingGuide, DevelopmentGuide, IntroToPandasGuide,
#                     BasicsOfODESystemGuide, BasicsOfDynSysImplementationGuide, BasicsOfReportingGuide, ResearchProjectGuidelines,
#                     InterimProjectGuidelines,IntroDynPyProjectGuidelines, BasicsOfReportComponentImplementationGuide,
#                     GithubSynchroGuide, IntroToCocalcGuide)
# from sympy import *
import datetime
import os
import shutil
from typing import Optional, Union

from pylatex import (  # Section, Subsection, Subsubsection, Itemize,  HorizontalSpace, Description, Marker
    Command,
    Document,
    NewPage,
    Package,
    Tabularx,
)
from pylatex.base_classes import Environment

# from pylatex.section import Paragraph, Chapter
from pylatex.utils import NoEscape  # italic,

from ..components.guides import en as guide_comp
from ..components.guides.development import en as development_comp
from ..components.guides.github import en as github_comp
from ..components.guides.pandas import en as pandas_comp
from ..components.guides.reporting import en as reporting_comp
from ..components.guides.systems import en as systems_comp
from ..components.mech import en as mech_compIntroDynPyProjectGuidlines
from ..components.ode import en as ode_comp_en
from ..components.ode import pl as ode_comp
from .guides import Guide, UsageOfDynamicSystemsGuide


class DynSysOverviewReport(UsageOfDynamicSystemsGuide):

    @property
    def default_reported_object(self):

        # from ...models.mechanics.tmac import SDOFWinchSystem
        from ...models.mechanics import ForcedSpringMassSystem as DynamicSystem

        return DynamicSystem()

    @property
    def _report_components(self):

        comp_list = [
            systems_comp.DynSysOverviewUsageComponent,
            systems_comp.DynamicSystemCallComponent,
            pandas_comp.NumericalAnalysisSimulationComponent,
            pandas_comp.AnalyticalSimulationComponent,
            # systems_comp.SimulationsComponent,
            # reporting_comp.SimulationReportComponent,
            systems_comp.DynSysCodeComponent,
            github_comp.IssuePreparationComponent,
            # systems_comp.DynamicSystemCheckerComponent,
        ]

        return comp_list


class ODESystemOverviewReport(UsageOfDynamicSystemsGuide):

    @property
    def _report_components(self):

        comp_list = [
            systems_comp.ODESystemOverviewUsageComponent,
            systems_comp.ODEInitCodeComponent,
            systems_comp.ODEGeneralSolutionComponent,
            systems_comp.ODESteadySolutionComponent,
            systems_comp.ODESystemRepresentationComponent,
            ode_comp_en.ODESystemExponentiationComponent,
            ode_comp_en.PredictionOfSteadySolutionComponent,
            ode_comp_en.HomoPredictionIntroComponent,
            ode_comp_en.MainPredictionComponent,
            ode_comp_en.RootsAnalysisComponent,
        ]

        return comp_list

    @property
    def default_reported_object(self):

        # from ...models.mechanics.tmac import SDOFWinchSystem
        from ...modf.odes.linear import LinearFirstOrder

        return LinearFirstOrder.from_reference_data()


class MSMethodOverviewReport(UsageOfDynamicSystemsGuide):
    @property
    def _report_components(self):
        comp_list = [
            ode_comp.ODESystemComponent,
            ode_comp.VariablesComponent,
            ode_comp.ODESystemCodeComponent,
            ode_comp.MSMCalculationsOrderComponent,
            ode_comp.PredictedSolutionComponent,
            ode_comp.DetailsOfPredictedSolutionComponent,
            ode_comp.ParticularDerivativesComponent,
            ode_comp_en.GoverningEqnsWithConstFuncsComponent,
            ode_comp_en.SecularFuncsComponent,
            ode_comp_en.SecularConditionsComponent,
            ode_comp_en.SolutionWithConstFuncsComponent,
            ode_comp.ZerothOrderApproximatedEqComponent,
            ode_comp.FirstOrderApproximatedEqComponent,
            ode_comp.SecularTermsEquationsComponent,
            ode_comp.ZerothOrderSolutionComponent,
        ]
        return comp_list

    @property
    def default_reported_object(self):
        from sympy import S, Symbol, sin

        from dynpy.models.mechanics.gears import EquivalentGearModel
        from dynpy.solvers.perturbational import MultiTimeScaleSolution

        sys = EquivalentGearModel()
        t = ivar = Symbol("t")
        z = sys.z
        delta = Symbol("delta", positive=True)
        eps = sys.eps
        nonlin_ode = MultiTimeScaleSolution(
            z.diff(t, 2) + z * (1 + delta * eps + eps * sin(t)),
            z,
            ivar=ivar,
            omega=S.One,
            order=3,
            eps=eps,
        )
        return nonlin_ode








# --- Thesis Card Generator --- #
from pathlib import Path
from datetime import datetime
import shutil, subprocess

try:
    from docx import Document
except Exception as e:
    raise RuntimeError("Zainstaluj zależność: pip install python-docx") from e


class ThesisCardMini:

    def __init__(self, template_path: Path | None = None, output_dir: Path | None = None, try_pdf: bool = True):
        if template_path is None:
            template_path = Path(__file__).resolve().parent / "documents" / "thesisCard.docx"
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Nie znaleziono szablonu: {self.template_path}")

        self.output_dir = Path(output_dir or (Path.cwd() / "build"))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.try_pdf = try_pdf

    @staticmethod
    def _canon(s: str) -> str:
        return (s or "").replace("\u00a0", " ").strip().lower().rstrip(":")

    def _set_label_in_paragraphs(self, doc, label: str, value: str) -> bool:

        lab = self._canon(label)
        for p in doc.paragraphs:
            if self._canon(p.text).startswith(lab):
                before = p.text.split(":", 1)[0] if ":" in p.text else label
                p.text = f"{before.strip()}: {value.strip()}"
                return True
        return False

    def _set_label_in_tables(self, doc, label: str, value: str) -> bool:

        lab = self._canon(label)
        for table in doc.tables:
            for row in table.rows:
                for j, cell in enumerate(row.cells):
                    txt = cell.text or ""
                    if self._canon(txt).startswith(lab):
                        if j + 1 < len(row.cells):
                            row.cells[j + 1].text = value.strip()
                        else:
                            before = txt.split(":", 1)[0] if ":" in txt else label
                            cell.text = f"{before.strip()}: {value.strip()}"
                        return True
        return False

    def _set_value_below_label(self, doc, label: str, value: str) -> bool:

        lab = self._canon(label)

        # akapity
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            if self._canon(p.text).startswith(lab):
                if i + 1 < len(paras):
                    paras[i + 1].text = value.strip()
                    return True
                return False

        # tabele
        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if self._canon(cell.text).startswith(lab):
                        rr = r_idx + 1
                        if rr < len(table.rows):
                            table.rows[rr].cells[c_idx].text = value.strip()
                            return True
                        if c_idx + 1 < len(row.cells):
                            row.cells[c_idx + 1].text = value.strip()
                            return True
                        before = (cell.text or label).split(":", 1)[0]
                        cell.text = f"{before.strip()}:\n{value.strip()}"
                        return True
        return False

    def _set_below_any_label(self, doc, labels: list[str], value: str) -> bool:

        for lab in labels:
            if self._set_value_below_label(doc, lab, value):
                return True
        return False

    def _set_tasks_after_label(self, doc, label: str, tasks: list[str]) -> bool:

        lab = self._canon(label)

        # akapity
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            if self._canon(p.text).startswith(lab):
                for k in range(5):
                    idx = i + 1 + k
                    if idx < len(paras):
                        txt = tasks[k].strip() if k < len(tasks) else ""
                        paras[idx].text = f"{k+1}. {txt}".rstrip()
                return True

        # tabele
        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if self._canon(cell.text).startswith(lab):
                        for k in range(5):
                            rr = r_idx + 1 + k
                            if rr < len(table.rows):
                                txt = tasks[k].strip() if k < len(tasks) else ""
                                table.rows[rr].cells[c_idx].text = f"{k+1}. {txt}".rstrip()
                        return True
        return False

    def _insert_date_above_label(self, doc, label: str | list[str], date_text: str, align: str = "center") -> bool:

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        labels = [label] if isinstance(label, str) else label
        labs = [self._canon(l) for l in labels]

        def _align_para(p, how: str):
            if how == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif how == "left":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        for p in doc.paragraphs:
            if any(l in self._canon(p.text) for l in labs):
                try:
                    new_p = p.insert_paragraph_before(date_text)
                except Exception:
                    p.text = f"{date_text}\n{p.text}"
                    _align_para(p, align)
                    return True
                _align_para(new_p, align)
                return True


        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if any(l in self._canon(cell.text) for l in labs):
                        # wiersz powyżej → ta sama kolumna
                        if r_idx > 0:
                            above = table.rows[r_idx - 1].cells[c_idx]
                            above.text = date_text
                            for para in above.paragraphs:
                                _align_para(para, align)
                            return True
                        # fallback: brak wiersza powyżej → prepend w tej samej komórce
                        if cell.text.strip():
                            cell.text = f"{date_text}\n{cell.text}"
                        else:
                            cell.text = date_text
                        for para in cell.paragraphs:
                            _align_para(para, align)
                        return True
        return False



    def _convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> bool:
        # 1) docx2pdf
        try:
            from docx2pdf import convert as docx2pdf_convert  # type: ignore
            docx2pdf_convert(str(docx_path), str(pdf_path))
            return pdf_path.exists()
        except Exception:
            pass
        # 2) LibreOffice
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            try:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                return pdf_path.exists()
            except Exception:
                return False
        return False



    def make_pdf(
        self,
        rodzaj: str = "niestacjonarne",
        stopien: str = "II",
        kierunek: str | None = None,
        specjalnosc: str | None = None,
        dyplomant: str | None = None,
        numer_albumu: str | None = None,
        email: str | None = None,
        prowadzacy: str | None = None,
        konsultant: str | None = None,
        temat: str | None = None,              # Temat POD etykietą (nowa linia)
        opis: str | None = None,               # Opis POD etykietą (nowa linia)
        jezyk: str | None = "Polski",          # Język opracowania po dwukropku
        zadania: list[str] | None = None,      # 1..5 zadań, opcjonalnie
        wstaw_date: bool = True,
        data_aktualna: str | None = None,      # np. "10.10.2025"; None -> dziś
        date_align: str = "center",
        signPath: str | Path | None = None,
        sign_x_mm: float | None = None,   # X od lewego-górnego rogu
        sign_y_mm: float | None = None,   # Y od lewego-górnego rogu
        sign_w_mm: float = 45.0,          # szerokość obrazka
        sign_h_mm: float | None = None,   # None = zachowaj proporcje                        
                                
                                                        # "center" (domyślnie) / "right" / "left"
    ) -> dict:
        from docx import Document
        doc = Document(str(self.template_path))

        # 1) standardowe pola 'po dwukropku'/ jeżeli zmieni się docelowy plik karty pracy, to może być konieczne dostosowanie tej części - na razie zostawiam jak jest.
        fields = {
            "Rodzaj studiów": rodzaj,
            "Stopień studiów": stopien,
            "Kierunek": kierunek,
            "Specjalność": specjalnosc,
            "Dyplomant": dyplomant,
            "Numer albumu": numer_albumu,
            "Adres e-mail": email,
            "Prowadzący": prowadzacy,
            "Konsultant": konsultant,
            "Język opracowania": jezyk,
        }

        not_found: list[str] = []
        for label, value in fields.items():
            if value is None or (isinstance(value, str) and value.strip() == ""):
                continue
            ok = self._set_label_in_paragraphs(doc, label, value)
            if not ok:
                ok = self._set_label_in_tables(doc, label, value)
            if not ok:
                not_found.append(label)


        if isinstance(temat, str) and temat.strip():
            ok_topic = self._set_below_any_label(
                doc,
                ["Temat pracy dyplomowej (nowy / zmieniony)*", "Temat pracy dyplomowej", "Temat pracy"],
                temat.strip(),
            )
            if not ok_topic:
                not_found.append("Temat pracy")


        if isinstance(opis, str) and opis.strip():
            opis_one_line = opis.replace("\r", " ").replace("\n", " ").strip()
            ok_opis = self._set_below_any_label(
                doc,
                ["Zwięzły opis celu pracy", "Zwięzły opis pracy", "Opis pracy"],
                opis_one_line,
            )
            if not ok_opis:
                not_found.append("Zwięzły opis celu pracy")


        if zadania:
            tasks = (zadania or [])[:5]
            ok_tasks = self._set_tasks_after_label(doc, "Główne zadania do wykonania", tasks)
            if not ok_tasks:
                not_found.append("Główne zadania do wykonania")

        if wstaw_date:
            today_raw = data_aktualna or datetime.now().strftime("%d.%m.%Y")
            trail = "\u00A0" * 37
            today = f"{today_raw}{trail}"

            ok_date = self._insert_date_above_label(
                doc,
                ["Data wydania", "Data wydania:", "data wydania"],
                today,
                align=date_align,
            )
            if not ok_date:
                not_found.append("Data wydania (wstawienie daty nad)")

        if not_found:
            raise ValueError("Nie znaleziono etykiet w szablonie: " + ", ".join(not_found))

        
        out_docx = self.output_dir / f"ThesisCard.docx"
        out_pdf  = self.output_dir / f"ThesisCard.pdf"

        doc.save(str(out_docx))
        made_pdf = self.try_pdf and self._convert_to_pdf(out_docx, out_pdf)

        # Add picture if path and picture are given
        if made_pdf and signPath:
            try:
                self.signPdf(
                    pdf_path=out_pdf,
                    sign_path=Path(signPath),
                    x_mm=sign_x_mm, y_mm=sign_y_mm,
                    width_mm=sign_w_mm, height_mm=sign_h_mm,
                )
            except Exception as e:
                return {"docx": str(out_docx), "pdf": str(out_pdf), "signature_error": str(e)}
    



    def signPdf(
        self,
        pdf_path: Path,
        sign_path: Path,
        *,
        x_mm: float | None,          
        y_mm: float | None,          
        width_mm: float = 45.0,
        height_mm: float | None = None,  # None => proporcje z obrazka
    ) -> bool:

        from pathlib import Path
        import io, os, time
        pdf_path = Path(pdf_path); sign_path = Path(sign_path)
        if not pdf_path.exists(): raise FileNotFoundError(f"PDF nie istnieje: {pdf_path}")
        if not sign_path.exists(): raise FileNotFoundError(f"Obraz podpisu nie istnieje: {sign_path}")
        if x_mm is None or y_mm is None:
            raise ValueError("Podaj x_mm i y_mm (mm od lewego-górnego rogu).")

        def mm2pt(mm: float) -> float: return mm * 72.0 / 25.4

        # --- Preferowany wariant: PyMuPDF (overlay + saveIncr) ---
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(pdf_path))
            page = doc[0]
            # rozmiar obrazka (do zachowania proporcji)
            try:
                pm = fitz.Pixmap(str(sign_path)); iw, ih = pm.width, pm.height; pm = None
            except Exception:
                iw, ih = 1, 1

            w_pt = mm2pt(width_mm)
            h_pt = mm2pt(height_mm) if height_mm else w_pt * (ih / iw)

            x0 = mm2pt(x_mm)
            y0 = mm2pt(y_mm)
            rect = fitz.Rect(x0, y0, x0 + w_pt, y0 + h_pt)

            page.insert_image(rect, filename=str(sign_path), keep_proportion=True, overlay=True)

            try:
                doc.saveIncr(); doc.close()
            except Exception:
                tmp = pdf_path.with_suffix(".signed.tmp.pdf")
                doc.save(str(tmp)); doc.close()
                try:
                    os.replace(tmp, pdf_path)
                finally:
                    if tmp.exists(): tmp.unlink(missing_ok=True)
            return True
        except Exception:
            # --- Fallback: ReportLab + PyPDF2 (też overlay) ---
            try:
                from PyPDF2 import PdfReader, PdfWriter
                from reportlab.pdfgen import canvas
                from reportlab.lib.utils import ImageReader
                import io

                reader = PdfReader(str(pdf_path))
                page0 = reader.pages[0]
                pw, ph = float(page0.mediabox.width), float(page0.mediabox.height)

                img = ImageReader(str(sign_path))
                iw, ih = img.getSize()
                w_pt = mm2pt(width_mm)
                h_pt = mm2pt(height_mm) if height_mm else w_pt * (ih / iw)

                # ReportLab ma (0,0) w lewym-DOLNYM rogu => konwersja Y:
                x = mm2pt(x_mm)
                y = ph - mm2pt(y_mm) - h_pt

                buf = io.BytesIO()
                c = canvas.Canvas(buf, pagesize=(pw, ph))
                c.drawImage(img, x, y, width=w_pt, height=h_pt, preserveAspectRatio=True, mask='auto')
                c.save(); buf.seek(0)

                overlay = PdfReader(buf).pages[0]
                page0.merge_page(overlay)

                writer = PdfWriter()
                for i, p in enumerate(reader.pages):
                    writer.add_page(page0 if i == 0 else p)

                tmp = pdf_path.with_suffix(".signed.tmp.pdf")
                with open(tmp, "wb") as f: writer.write(f)
                try:
                    os.replace(tmp, pdf_path)
                finally:
                    if os.path.exists(tmp): os.remove(tmp)
                return True
            except Exception:
                return False